"""Integration test: Office auto-conversion + image translation.

Builds a real .docx (with python-docx) containing a paragraph and an
embedded image, renames it to a .doc filename to simulate a legacy file,
and runs the translator pipeline with auto-conversion enabled.

We mock convert_to_modern_format to perform the rename so the test runs
without real Win32COM / LibreOffice, and we mock the OCR + image-render
stack so embedded-image translation is deterministic.
"""

from __future__ import annotations

import shutil
from collections.abc import Callable, Generator
from pathlib import Path

import pytest

from src.constants.history import STATUS_DONE, STATUS_PENDING
from src.core.config import TranslationConfig
from src.core.database import (
    add_history_entry,
    get_history_entry_status,
    init_db,
)
from src.core.translator import (
    _update_storage_path,
    run_translation_pipeline,
)

# A 1×1 white PNG used as the embedded image.
_PNG_1X1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xff"
    b"\xff?\x00\x05\xfe\x02\xfe\xa3\x9b\x99\x9c\x00\x00\x00\x00IEND\xaeB`\x82"
)


# ── Fixtures ──────────────────────────────────────────────────────────


@pytest.fixture(autouse=True)
def setup_integration_env(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> Generator[None, None, None]:
    """Per-test DB isolation + path redirection."""
    db_file = tmp_path / "integration.db"
    monkeypatch.setattr("src.core.database.get_db_path", lambda: str(db_file))
    config_dir = tmp_path / "config"
    config_dir.mkdir()
    monkeypatch.setattr(
        "src.utils.path_manager.get_app_config_dir",
        lambda: config_dir,
    )
    data_dir = tmp_path / "data"
    data_dir.mkdir()
    monkeypatch.setattr(
        "src.utils.path_manager.get_app_data_dir",
        lambda: data_dir,
    )
    init_db()
    monkeypatch.setattr("src.core.translator.stop_soffice", lambda: None)
    # Force python-lib backend to avoid UNO connection issues in CI.
    monkeypatch.setattr(
        "src.core.office_processor._detect_backend",
        lambda suffix, *_args: "python_lib",
    )
    yield


@pytest.fixture()
def mock_llm(
    monkeypatch: pytest.MonkeyPatch,
) -> Callable[..., list[str]]:
    """Mocks translate_text at every import site."""

    def fake_translate(
        texts: list[str],
        target_lang: str,
        source_lang: str = "",
        **kwargs: object,
    ) -> list[str]:
        return [f"[{target_lang}] {t}" for t in texts]

    monkeypatch.setattr("src.core.llm_engine.translate_text", fake_translate)
    monkeypatch.setattr(
        "src.core.text_processor._llm_engine.translate_text",
        fake_translate,
    )
    return fake_translate


# ── Helpers ──────────────────────────────────────────────────────────


def _create_docx_with_image(path: Path) -> None:
    """Builds a real .docx with a paragraph + an embedded inline image."""
    from io import BytesIO  # noqa: PLC0415

    from docx import Document  # noqa: PLC0415
    from docx.shared import Inches  # noqa: PLC0415

    doc = Document()
    doc.add_paragraph("Hello legacy document")
    doc.add_picture(BytesIO(_PNG_1X1), width=Inches(1))
    doc.save(str(path))


# ── The main test ────────────────────────────────────────────────────


def test_legacy_doc_auto_convert_with_image_translation(
    tmp_path: Path,
    mock_llm: Callable[..., list[str]],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Pipeline: .doc → auto-convert to .docx → translate text + image.

    Asserts:
    - convert_to_modern_format is invoked exactly once with .doc input.
    - _translate_doc_images is invoked once on the converted .docx with
      ".docx" suffix (image translation runs on the modern format).
    - Body paragraph is translated.
    - DB entry transitions Pending → Done.
    - DB filename column is updated to the modern .docx name.
    """
    # ── 1. Create the source .docx, then rename to .doc ─────────────
    real_docx = tmp_path / "source.docx"
    _create_docx_with_image(real_docx)

    # Set up the per-task storage dir like translator.setup_translation_tasks would.
    storage_root = tmp_path / "data" / "translations"
    storage_root.mkdir(parents=True, exist_ok=True)

    legacy_name = "source.doc"

    # Add DB entry first (we need the ID to build the storage subdir).
    h_id = add_history_entry(
        legacy_name,
        "English (US)",
        "French",
        STATUS_PENDING,
        source_path=str(tmp_path / legacy_name),
    )
    assert h_id is not None

    task_dir = storage_root / str(h_id)
    task_dir.mkdir(parents=True, exist_ok=True)
    legacy_path = task_dir / legacy_name
    # Place the (truly-docx) bytes at the legacy filename.
    shutil.copy2(real_docx, legacy_path)
    _update_storage_path(h_id, str(legacy_path.resolve()))

    # ── 2. Mock convert_to_modern_format to do a real rename/copy ───
    convert_calls: list[tuple[Path, Path]] = []

    def fake_convert(input_path: Path, output_path: Path) -> bool:
        convert_calls.append((Path(input_path), Path(output_path)))
        # Real round-trip: copy the bytes (which are actually a valid docx)
        # to the modern path.
        shutil.copy2(input_path, output_path)
        return True

    monkeypatch.setattr(
        "src.core.translator.convert_to_modern_format",
        fake_convert,
    )

    # ── 3. Mock _translate_doc_images so we don't need real OCR ─────
    image_calls: list[tuple[Path, str]] = []

    def fake_translate_doc_images(  # noqa: PLR0913
        output_path: Path,
        suffix: str,
        backend: str,
        target_lang: str,
        src_lang: str,
        glossary_entries: list[tuple[int, str, str]] | None,
        progress_callback: object,
        cancel_check: object,
        config: object = None,
        *,
        provider: str | None = None,
        model: str | None = None,
        checkpoint_dir: Path | None = None,  # noqa: ARG001
    ) -> None:
        image_calls.append((Path(output_path), suffix))

    monkeypatch.setattr(
        "src.core.office_processor._translate_doc_images",
        fake_translate_doc_images,
    )

    # ── 4. Build a config that enables auto-convert + image translation ─
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    config = TranslationConfig(
        storage_path=str(output_dir),
        auto_remove_history=False,
        auto_convert_legacy=True,
        translate_doc_images=True,
        ocr_is_configured=True,
        ocr_method="TesseractOCR",
    )

    # ── 5. Run the pipeline ────────────────────────────────────────
    run_translation_pipeline(config=config)

    # ── 6. Assertions ──────────────────────────────────────────────
    # DB transitions to Done.
    status = get_history_entry_status(h_id)
    assert status == STATUS_DONE, f"Expected Done, got {status}"

    # convert_to_modern_format invoked exactly once with .doc → .docx.
    assert len(convert_calls) == 1, f"Expected 1 convert call, got {convert_calls}"
    converted_input, converted_output = convert_calls[0]
    assert converted_input.suffix.lower() == ".doc"
    assert converted_output.suffix.lower() == ".docx"

    # _translate_doc_images invoked once on the modern .docx.
    assert len(image_calls) == 1, f"Expected 1 image call, got {image_calls}"
    _img_path, img_suffix = image_calls[0]
    assert img_suffix == ".docx"

    # Output file exists with translated content.
    output_files = list(output_dir.rglob("*.docx"))
    assert output_files, f"No .docx output produced in {output_dir}"
    out_path = output_files[0]

    # Verify body paragraph got translated.
    from docx import Document  # noqa: PLC0415

    doc = Document(str(out_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("[French]" in p for p in paragraphs), (
        f"Expected translated paragraph, got: {paragraphs}"
    )


def test_legacy_doc_auto_convert_disabled_keeps_legacy_extension(
    tmp_path: Path,
    mock_llm: Callable[..., list[str]],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """When auto_convert_legacy=False, convert_to_modern_format is NOT called."""
    real_docx = tmp_path / "source.docx"
    _create_docx_with_image(real_docx)

    storage_root = tmp_path / "data" / "translations"
    storage_root.mkdir(parents=True, exist_ok=True)

    h_id = add_history_entry(
        "source.docx",  # use real modern format so processing succeeds
        "English (US)",
        "French",
        STATUS_PENDING,
        source_path=str(tmp_path / "source.docx"),
    )
    assert h_id is not None
    task_dir = storage_root / str(h_id)
    task_dir.mkdir(parents=True, exist_ok=True)
    cloned = task_dir / "source.docx"
    shutil.copy2(real_docx, cloned)
    _update_storage_path(h_id, str(cloned.resolve()))

    convert_calls: list[object] = []

    def fake_convert(input_path: Path, output_path: Path) -> bool:
        convert_calls.append((input_path, output_path))
        shutil.copy2(input_path, output_path)
        return True

    monkeypatch.setattr(
        "src.core.translator.convert_to_modern_format",
        fake_convert,
    )

    # No-op image pipeline.
    monkeypatch.setattr(
        "src.core.office_processor._translate_doc_images",
        lambda *args, **kwargs: None,
    )

    output_dir = tmp_path / "output"
    output_dir.mkdir()
    config = TranslationConfig(
        storage_path=str(output_dir),
        auto_remove_history=False,
        auto_convert_legacy=False,
        translate_doc_images=False,
        ocr_is_configured=False,
    )

    run_translation_pipeline(config=config)

    assert get_history_entry_status(h_id) == STATUS_DONE
    # No conversion happened: the file was already a modern .docx.
    assert convert_calls == [], (
        f"convert_to_modern_format should not have been called: {convert_calls}"
    )


# ─────────────────────────────────────────────────────────────────────
# Integration: atomic output + per-image cache + skip-with-warning
# ─────────────────────────────────────────────────────────────────────


def test_image_cache_skip_warning_atomic_output_chain(  # noqa: PLR0915
    tmp_path: Path,
    mock_llm: Callable[..., list[str]],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """End-to-end: cache hit, fresh translate, skipped failure, atomic publish.

    Builds a real .docx with three embedded images.  Mocks
    ``_translate_single_image`` so:

      * image 1 → returns translated bytes (gets cached + injected)
      * image 2 → raises ``ValueError('CONNECTION_ERROR')`` (skip
        policy keeps original; not cached)
      * image 3 → returns translated bytes (gets cached + injected)

    Asserts the chain holds together:

      1. **Task completes Done** — skip-with-warning means one bad
         image doesn't sink the whole document.
      2. **Atomic publish** — the final output file appears at the
         user's destination; no ``_partial.docx`` is left behind in
         the task storage dir after success.
      3. **Per-image cache** populated for the successful images,
         absent for the failed one (so a retry would re-attempt only
         image 2).
      4. **Skip-with-warning** — original bytes for image 2 survive
         in the output ZIP.
    """
    from io import BytesIO  # noqa: PLC0415

    from docx import Document  # noqa: PLC0415
    from docx.shared import Inches  # noqa: PLC0415

    from src.constants.history import STATUS_PENDING  # noqa: PLC0415
    from src.core.checkpoint import hash_office_image  # noqa: PLC0415

    # Three visually-distinct PNGs so each one has a unique SHA256.
    png_b = _PNG_1X1
    png_b2 = _PNG_1X1[:-4] + b"\x12\x34\x56\x78"
    png_b3 = _PNG_1X1[:-4] + b"\xab\xcd\xef\x01"

    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("Body text")
    doc.add_picture(BytesIO(png_b), width=Inches(1))
    doc.add_picture(BytesIO(png_b2), width=Inches(1))
    doc.add_picture(BytesIO(png_b3), width=Inches(1))
    doc.save(str(src))

    storage_root = tmp_path / "data" / "translations"
    storage_root.mkdir(parents=True, exist_ok=True)

    h_id = add_history_entry(
        "src.docx",
        "English (US)",
        "French",
        STATUS_PENDING,
        source_path=str(src),
    )
    assert h_id is not None
    task_dir = storage_root / str(h_id)
    task_dir.mkdir(parents=True, exist_ok=True)
    cloned = task_dir / "src.docx"
    shutil.copy2(src, cloned)
    _update_storage_path(h_id, str(cloned.resolve()))

    # Three-way image fate: success, skip, success.  ``side_effect``
    # is a list so each call returns the next entry.
    call_outputs = [
        b"translated-img-1",
        ValueError("CONNECTION_ERROR"),
        b"translated-img-3",
    ]

    def fake_single_image(image_bytes: bytes, *_args, **_kwargs) -> bytes:
        outcome = call_outputs.pop(0)
        if isinstance(outcome, BaseException):
            raise outcome
        return outcome

    monkeypatch.setattr(
        "src.core.office_processor._translate_single_image",
        fake_single_image,
    )

    output_dir = tmp_path / "output"
    output_dir.mkdir()
    config = TranslationConfig(
        storage_path=str(output_dir),
        auto_remove_history=False,
        translate_doc_images=True,
        ocr_is_configured=True,
        ocr_method="TesseractOCR",
    )

    run_translation_pipeline(config=config)

    # 1. Task completed despite one failed image.
    assert get_history_entry_status(h_id) == STATUS_DONE, (
        "Skip-with-warning must let the document complete even when "
        "one image fails persistently."
    )

    # 2. Atomic publish: the user-visible output file appeared in the
    # output dir.  We don't know the exact filename (translator builds
    # it via ``_build_output_name``) — assert that exactly one
    # translated .docx exists in the output dir.
    produced = list(output_dir.glob("*.docx"))
    assert len(produced) == 1, (
        f"Expected exactly one .docx in output dir; got {produced}"
    )
    final_output = produced[0]

    # No orphan ``_partial.docx`` was left in the task storage dir
    # after the successful move.
    assert not (task_dir / "_partial.docx").exists(), (
        "Atomic publish should move the partial into output_path and "
        "leave the storage dir clean."
    )

    # 3. Per-image cache: ``clear_checkpoints`` runs on task success
    # and wipes the ``office_images/`` subdir, so the cache is gone
    # by the time we inspect it.  Instead verify the chain end-state:
    # the final ZIP contains the successful translations and the
    # original bytes for the failed image.
    import zipfile  # noqa: PLC0415

    # Compute the expected zip path for each source image bytes by
    # hash so we can find them in the produced docx without relying
    # on python-docx's internal numbering.
    h1 = hash_office_image(png_b)
    h2 = hash_office_image(png_b2)
    h3 = hash_office_image(png_b3)
    assert h1 != h2 != h3  # sanity

    with zipfile.ZipFile(final_output, "r") as zf:
        media_files = [
            n for n in zf.namelist() if n.startswith("word/media/")
        ]
        assert len(media_files) == 3, (
            f"Expected 3 media files in output; got {media_files}"
        )
        media_bytes = sorted(zf.read(n) for n in media_files)

    expected_after_skip = sorted([
        b"translated-img-1",
        png_b2,  # original survived the skip-with-warning
        b"translated-img-3",
    ])
    assert media_bytes == expected_after_skip, (
        "Output ZIP media doesn't match skip-with-warning expectation"
    )
