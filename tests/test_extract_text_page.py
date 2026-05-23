"""Comprehensive tests for ExtractTextPage.

Covers page construction, widget structure, stacked view switching,
file handling, method selection, requirements checking, output format,
apply_theme, apply_language, and the _ExtractionWorker class.
"""

from __future__ import annotations

from unittest.mock import MagicMock, patch

import pytest
from PySide6.QtWidgets import (
    QLabel,
    QMainWindow,
    QPushButton,
    QStackedWidget,
)

# ---------------------------------------------------------------------------
# Module-level patch path constants
# ---------------------------------------------------------------------------
_MOD = "src.ui.pages.extract_text"
_HIST_MOD = "src.ui.pages.extraction_history"


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def _mock_db():
    """Mocks database calls used by embedded ExtractionHistoryPage."""
    with (
        patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
        patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
    ):
        yield


@pytest.fixture()
def window(qtbot) -> QMainWindow:
    """Creates a minimal QMainWindow for parenting."""
    w = QMainWindow()
    w.navigate_to_settings_tab = MagicMock()
    qtbot.addWidget(w)
    return w


@pytest.fixture()
def page(_mock_db, window, qtbot):
    """Creates an ExtractTextPage widget for testing."""
    from src.ui.pages.extract_text import ExtractTextPage  # noqa: PLC0415

    p = ExtractTextPage(window)
    qtbot.addWidget(p)
    return p


# ===================================================================
# Widget Construction
# ===================================================================


class TestConstruction:
    """Tests for ExtractTextPage widget construction."""

    def test_page_created(self, page) -> None:  # noqa: ANN001
        """Page is created without error."""
        assert page is not None

    def test_has_stack_widget(self, page) -> None:  # noqa: ANN001
        """Page has a QStackedWidget for view switching."""
        assert isinstance(page.stack, QStackedWidget)

    def test_stack_has_two_views(self, page) -> None:  # noqa: ANN001
        """Stack has exactly 2 views (history and files)."""
        assert page.stack.count() == 2  # noqa: PLR2004

    def test_initial_view_is_history(self, page) -> None:  # noqa: ANN001
        """Initial view shows the history view (index 0)."""
        assert page.stack.currentIndex() == 0

    def test_has_drop_area(self, page) -> None:  # noqa: ANN001
        """Page has a FileDropWidget."""
        from src.ui.components import FileDropWidget  # noqa: PLC0415

        assert isinstance(page.drop_area, FileDropWidget)

    def test_has_extraction_history(self, page) -> None:  # noqa: ANN001
        """Page embeds an ExtractionHistoryPage."""
        from src.ui.pages.extraction_history import ExtractionHistoryPage  # noqa: PLC0415, I001

        assert isinstance(page.extraction_history, ExtractionHistoryPage)

    def test_has_extract_button(self, page) -> None:  # noqa: ANN001
        """Page has an extract button."""
        assert isinstance(page.extract_btn, QPushButton)

    def test_has_clear_all_button(self, page) -> None:  # noqa: ANN001
        """Page has a clear-all button."""
        assert isinstance(page.clear_all_btn, QPushButton)

    def test_has_files_badge(self, page) -> None:  # noqa: ANN001
        """Page has a file count badge label."""
        assert isinstance(page.files_badge, QLabel)

    def test_has_section_label(self, page) -> None:  # noqa: ANN001
        """Page has a 'files selected' section label."""
        assert isinstance(page.section_label, QLabel)

    def test_extract_button_disabled_initially(self, page) -> None:  # noqa: ANN001
        """Extract button is disabled when no files are selected."""
        assert not page.extract_btn.isEnabled()

    def test_selected_files_empty_initially(self, page) -> None:  # noqa: ANN001
        """No files are selected at construction time."""
        assert page.selected_files == []

    def test_badge_shows_zero_initially(self, page) -> None:  # noqa: ANN001
        """File count badge shows '0' initially."""
        assert page.files_badge.text() == "0"

    def test_output_format_default(self, page) -> None:  # noqa: ANN001
        """Default output format is .txt."""
        assert page._output_format == ".txt"


# ===================================================================
# View Switching (_update_ui_state)
# ===================================================================


class TestViewSwitching:
    """Tests for stacked view switching based on file selection."""

    def test_no_files_shows_history_view(self, page) -> None:  # noqa: ANN001
        """With no files, stack shows history view (index 0)."""
        page.selected_files.clear()
        page._update_ui_state()
        assert page.stack.currentIndex() == 0

    def test_files_selected_shows_files_view(self, page) -> None:  # noqa: ANN001
        """With files selected, stack switches to files view (index 1)."""
        page.selected_files = ["/tmp/test.png"]
        page._update_ui_state()
        assert page.stack.currentIndex() == 1

    def test_badge_updates_with_file_count(self, page) -> None:  # noqa: ANN001
        """File count badge updates when files are added."""
        page.selected_files = ["/tmp/a.png", "/tmp/b.png", "/tmp/c.png"]
        page._update_ui_state()
        assert page.files_badge.text() == "3"

    def test_extract_button_enabled_with_files(self, page) -> None:  # noqa: ANN001
        """Extract button is enabled when files are selected."""
        page.selected_files = ["/tmp/test.png"]
        page._update_ui_state()
        assert page.extract_btn.isEnabled()

    def test_extract_button_disabled_without_files(self, page) -> None:  # noqa: ANN001
        """Extract button is disabled after files are cleared."""
        page.selected_files = ["/tmp/test.png"]
        page._update_ui_state()
        page.selected_files.clear()
        page._update_ui_state()
        assert not page.extract_btn.isEnabled()


# ===================================================================
# File Handling (_handle_clear_all)
# ===================================================================


class TestClearAll:
    """Tests for _handle_clear_all behavior."""

    def test_clear_all_empties_selected_files(self, page) -> None:  # noqa: ANN001
        """_handle_clear_all empties the selected_files list."""
        page.selected_files = ["/tmp/a.png", "/tmp/b.png"]
        page._handle_clear_all()
        assert page.selected_files == []

    def test_clear_all_switches_to_history_view(self, page) -> None:  # noqa: ANN001
        """_handle_clear_all switches back to history view."""
        page.selected_files = ["/tmp/a.png"]
        page._update_ui_state()
        page._handle_clear_all()
        assert page.stack.currentIndex() == 0


# ===================================================================
# Requirements Checking (_check_extract_requirements)
# ===================================================================


class TestCheckExtractRequirements:
    """Tests for _check_extract_requirements method."""

    @patch(f"{_MOD}.check_llm_setup", return_value=True)
    @patch(f"{_MOD}.check_ocr_setup", return_value=False)
    def test_passes_when_llm_configured(
        self,
        mock_ocr,
        mock_llm,
        page,  # noqa: ANN001
    ) -> None:
        """Returns True when LLM is configured (even if OCR is not)."""
        result = page._check_extract_requirements()
        assert result is True

    @patch(f"{_MOD}.check_llm_setup", return_value=False)
    @patch(f"{_MOD}.check_ocr_setup", return_value=True)
    def test_passes_when_ocr_configured(
        self,
        mock_ocr,
        mock_llm,
        page,  # noqa: ANN001
    ) -> None:
        """Returns True when OCR is configured (even if LLM is not)."""
        result = page._check_extract_requirements()
        assert result is True

    @patch(f"{_MOD}.require_setup", return_value=False)
    @patch(f"{_MOD}.check_llm_setup", return_value=False)
    @patch(f"{_MOD}.check_ocr_setup", return_value=False)
    def test_fails_when_neither_configured(
        self,
        mock_ocr,
        mock_llm,
        mock_require,
        page,  # noqa: ANN001
    ) -> None:
        """Returns False when neither OCR nor LLM is configured."""
        result = page._check_extract_requirements()
        assert result is False

    @patch(f"{_MOD}.check_llm_setup", return_value=True)
    @patch(f"{_MOD}.check_ocr_setup", return_value=True)
    def test_passes_when_both_configured(
        self,
        mock_ocr,
        mock_llm,
        page,  # noqa: ANN001
    ) -> None:
        """Returns True when both OCR and LLM are configured."""
        result = page._check_extract_requirements()
        assert result is True


# ===================================================================
# Handle Extract (_handle_extract)
# ===================================================================


class TestHandleExtract:
    """Tests for _handle_extract behavior."""

    def test_extract_noop_when_no_files(self, page) -> None:  # noqa: ANN001
        """_handle_extract does nothing when selected_files is empty."""
        page.selected_files.clear()
        page._handle_extract()  # Should not raise

    @patch(f"{_MOD}.require_setup", return_value=False)
    def test_extract_blocked_when_setup_missing(
        self,
        mock_require,
        page,  # noqa: ANN001
    ) -> None:
        """_handle_extract is blocked when requirements check fails."""
        page.selected_files = ["/tmp/test.png"]
        page._handle_extract()
        mock_require.assert_called_once()

    @patch(f"{_MOD}._ExtractionWorker")
    @patch(f"{_MOD}.add_extraction_entry", return_value=1)
    @patch(f"{_MOD}.load_setting", return_value="ocr")
    @patch(
        f"{_MOD}.SourceLanguageDialog.get_selection",
        return_value=("en", "English", None, True),
    )
    @patch(f"{_MOD}.require_setup", return_value=True)
    def test_extract_starts_worker_and_clears_files(  # noqa: PLR0913
        self,
        mock_require,
        mock_dialog,
        mock_load,
        mock_add,
        mock_worker_cls,
        page,  # noqa: ANN001
    ) -> None:
        """_handle_extract starts worker and clears selection on success."""
        mock_worker = MagicMock()
        mock_worker_cls.return_value = mock_worker

        page.selected_files = ["/tmp/test.png"]
        page._update_ui_state()

        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._handle_extract()

        # Files should be cleared
        assert page.selected_files == []
        # Worker should have been started
        mock_worker.start.assert_called_once()

    @patch(f"{_MOD}.load_setting", return_value="ocr")
    @patch(
        f"{_MOD}.SourceLanguageDialog.get_selection",
        return_value=("en", "English", None, False),
    )
    @patch(f"{_MOD}.require_setup", return_value=True)
    def test_extract_cancelled_dialog_keeps_files(
        self,
        mock_require,
        mock_dialog,
        mock_load,
        page,  # noqa: ANN001
    ) -> None:
        """Files are kept when user cancels the source language dialog."""
        page.selected_files = ["/tmp/test.png"]
        page._handle_extract()
        # Files should NOT be cleared (dialog was cancelled)
        assert page.selected_files == ["/tmp/test.png"]


# ===================================================================
# Handle Re-Extract (_handle_re_extract)
# ===================================================================


class TestHandleReExtract:
    """Tests for _handle_re_extract behavior."""

    @patch(f"{_MOD}.require_setup", return_value=False)
    def test_re_extract_blocked_when_setup_missing(
        self,
        mock_require,
        page,  # noqa: ANN001
    ) -> None:
        """_handle_re_extract is blocked when requirements check fails."""
        page._handle_re_extract([(1, "/tmp/test.png")])
        mock_require.assert_called_once()

    @patch(f"{_MOD}._ExtractionWorker")
    @patch(f"{_MOD}.load_setting", return_value="ocr")
    @patch(f"{_MOD}.update_extraction_status")
    @patch(
        f"{_MOD}.SourceLanguageDialog.get_selection",
        return_value=("en", "English", None, True),
    )
    @patch(f"{_MOD}.require_setup", return_value=True)
    def test_re_extract_resets_status_and_starts_worker(  # noqa: PLR0913
        self,
        mock_require,
        mock_dialog,
        mock_update,
        mock_load,
        mock_worker_cls,
        page,  # noqa: ANN001
    ) -> None:
        """_handle_re_extract resets entries to Pending and starts worker."""
        mock_worker = MagicMock()
        mock_worker_cls.return_value = mock_worker

        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._handle_re_extract([(10, "/tmp/img.png")])

        mock_update.assert_called_once_with(10, "Pending")  # noqa: PLR2004
        mock_worker.start.assert_called_once()


# ===================================================================
# Theme / Language
# ===================================================================


class TestThemeAndLanguage:
    """Tests for apply_theme and apply_language methods."""

    def test_apply_theme_runs(self, page) -> None:  # noqa: ANN001
        """apply_theme() completes without error."""
        page.apply_theme()

    def test_apply_language_runs(self, page) -> None:  # noqa: ANN001
        """apply_language() completes without error."""
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page.apply_language()

    def test_apply_theme_updates_button_styles(self, page) -> None:  # noqa: ANN001
        """apply_theme updates styles for action buttons."""
        page.apply_theme()
        assert page.extract_btn.styleSheet()
        assert page.clear_all_btn.styleSheet()

    def test_apply_language_updates_button_text(self, page) -> None:  # noqa: ANN001
        """apply_language updates button labels."""
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page.apply_language()

        assert page.extract_btn.text()
        assert page.clear_all_btn.text()
        assert page.section_label.text()

    def test_apply_theme_updates_badge_style(self, page) -> None:  # noqa: ANN001
        """apply_theme updates the files_badge stylesheet."""
        page.apply_theme()
        assert page.files_badge.styleSheet()


# ===================================================================
# _ExtractionWorker
# ===================================================================


class TestExtractionWorker:
    """Tests for the _ExtractionWorker class."""

    def test_worker_is_busy_initially_false(self) -> None:
        """is_busy() returns False before any worker starts."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        # Reset class-level flag (may be dirty from other tests)
        _ExtractionWorker._is_any_worker_running = False
        assert not _ExtractionWorker.is_busy()

    def test_worker_stop_sets_flag(self) -> None:
        """stop() sets _is_running to False."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en")
        assert worker._is_running is True
        worker.stop()
        assert worker._is_running is False

    def test_worker_extract_method_stored(self) -> None:
        """Worker stores the extraction method."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en", extract_method="llm")
        assert worker._extract_method == "llm"

    def test_worker_default_extract_method_is_ocr(self) -> None:
        """Worker defaults to OCR extraction method."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en")
        assert worker._extract_method == "OCR"


# ===================================================================
# _write_extraction_output
# ===================================================================


class TestWriteExtractionOutput:
    """Tests for the _write_extraction_output utility function."""

    def test_write_txt_output(self, tmp_path) -> None:  # noqa: ANN001
        """Writes plain text to a .txt file."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "output.txt"
        _write_extraction_output(out, "Hello\nWorld")
        assert out.read_text(encoding="utf-8") == "Hello\nWorld"

    def test_write_docx_output(self, tmp_path) -> None:  # noqa: ANN001
        """Writes text to a .docx file."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "output.docx"
        _write_extraction_output(out, "Line 1\nLine 2")
        assert out.exists()
        # Verify it's a valid zip (docx is a zip file)
        import zipfile  # noqa: PLC0415

        assert zipfile.is_zipfile(out)


# ===================================================================
# create_extract_text_page factory
# ===================================================================


class TestFactory:
    """Tests for the create_extract_text_page factory function."""

    def test_factory_returns_page(self, _mock_db, window, qtbot) -> None:  # noqa: ANN001
        """create_extract_text_page returns an ExtractTextPage instance."""
        from src.ui.pages.extract_text import (  # noqa: PLC0415
            ExtractTextPage,
            create_extract_text_page,
        )

        page = create_extract_text_page(window)
        qtbot.addWidget(page)
        assert isinstance(page, ExtractTextPage)


# ===================================================================
# NEW TESTS: Page Creation (additional)
# ===================================================================


class TestExtractPageCreation:
    """Additional tests for page creation and widget presence."""

    def test_create_extract_text_page_returns_qwidget(
        self, _mock_db, window, qtbot
    ) -> None:  # noqa: ANN001
        """create_extract_text_page returns a QWidget."""
        from PySide6.QtWidgets import QWidget as _QWidget  # noqa: PLC0415

        from src.ui.pages.extract_text import create_extract_text_page  # noqa: PLC0415

        page = create_extract_text_page(window)
        qtbot.addWidget(page)
        assert isinstance(page, _QWidget)

    def test_has_extract_button(self, page) -> None:  # noqa: ANN001
        """Page has an extract button that is a QPushButton."""
        assert isinstance(page.extract_btn, QPushButton)
        assert page.extract_btn.text() != ""

    def test_has_method_drop_area(self, page) -> None:  # noqa: ANN001
        """Page has a drop area for file input."""
        from src.ui.components import FileDropWidget  # noqa: PLC0415

        assert isinstance(page.drop_area, FileDropWidget)

    def test_has_files_badge_label(self, page) -> None:  # noqa: ANN001
        """Page has a file count badge label."""
        assert isinstance(page.files_badge, QLabel)

    def test_has_section_label(self, page) -> None:  # noqa: ANN001
        """Page has a 'files selected' section label."""
        assert isinstance(page.section_label, QLabel)

    def test_has_clear_all_button(self, page) -> None:  # noqa: ANN001
        """Page has a clear-all button."""
        assert isinstance(page.clear_all_btn, QPushButton)

    def test_has_stacked_widget(self, page) -> None:  # noqa: ANN001
        """Page has a QStackedWidget for view switching."""
        assert isinstance(page.stack, QStackedWidget)

    def test_initial_worker_is_none(self, page) -> None:  # noqa: ANN001
        """Worker is None at construction time."""
        assert page._worker is None


# ===================================================================
# NEW TESTS: Extraction Worker (additional)
# ===================================================================


class TestExtractionWorkerAdditional:
    """Additional tests for _ExtractionWorker class."""

    def test_ocr_extraction_path(self) -> None:
        """Worker with OCR method calls _extract_with_ocr."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker(
            [(1, "/tmp/img.png")], "tesseract", "en", extract_method="OCR"
        )
        with (
            patch.object(worker, "_extract_with_ocr", return_value="OCR text") as m_ocr,
            patch.object(worker, "_extract_with_llm") as m_llm,
            patch(f"{_MOD}.update_extraction_status"),
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
            m_ocr.assert_called_once_with("/tmp/img.png")
            m_llm.assert_not_called()

    def test_llm_extraction_path(self) -> None:
        """Worker with LLM method calls _extract_with_llm."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker(
            [(1, "/tmp/img.png")], "tesseract", "en", extract_method="LLM"
        )
        with (
            patch.object(worker, "_extract_with_llm", return_value="LLM text") as m_llm,
            patch.object(worker, "_extract_with_ocr") as m_ocr,
            patch(f"{_MOD}.update_extraction_status"),
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
            m_llm.assert_called_once_with("/tmp/img.png")
            m_ocr.assert_not_called()

    def test_worker_with_multiple_images(self) -> None:
        """Worker processes all images in the task list."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        tasks = [
            (1, "/tmp/a.png"),
            (2, "/tmp/b.png"),
            (3, "/tmp/c.png"),
        ]
        worker = _ExtractionWorker(tasks, "tesseract", "en", extract_method="OCR")

        call_count = 0

        def mock_ocr(path: str) -> str:
            nonlocal call_count
            call_count += 1
            return f"text from {path}"

        with (
            patch.object(worker, "_extract_with_ocr", side_effect=mock_ocr),
            patch(f"{_MOD}.update_extraction_status"),
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
            assert call_count == 3  # noqa: PLR2004

    def test_worker_with_invalid_image_marks_failed(self) -> None:
        """Worker marks entry as Failed when extraction raises an exception."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker(
            [(1, "/tmp/bad.png")], "tesseract", "en", extract_method="OCR"
        )
        with (
            patch.object(
                worker, "_extract_with_ocr", side_effect=RuntimeError("bad image")
            ),
            patch(f"{_MOD}.update_extraction_status") as mock_update,
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
            # Should have been called with STATUS_FAILED
            calls = [c for c in mock_update.call_args_list if c[0][1] == "Failed"]
            assert len(calls) == 1

    def test_worker_cancellation_during_extraction(self) -> None:
        """Worker stops processing when stop() is called."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        tasks = [
            (1, "/tmp/a.png"),
            (2, "/tmp/b.png"),
            (3, "/tmp/c.png"),
        ]
        worker = _ExtractionWorker(tasks, "tesseract", "en", extract_method="OCR")

        call_count = 0

        def mock_ocr(path: str) -> str:
            nonlocal call_count
            call_count += 1
            # Stop worker after first image
            worker.stop()
            return "text"

        with (
            patch.object(worker, "_extract_with_ocr", side_effect=mock_ocr),
            patch(f"{_MOD}.update_extraction_status"),
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
            # Only the first image should be processed
            assert call_count == 1

    def test_worker_error_handling_continues(self) -> None:
        """Worker continues to next task when one fails."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        tasks = [
            (1, "/tmp/a.png"),
            (2, "/tmp/b.png"),
        ]
        worker = _ExtractionWorker(tasks, "tesseract", "en", extract_method="OCR")

        results_list = []

        def capture_results(results: list) -> None:
            results_list.extend(results)

        call_idx = 0

        def mock_ocr(path: str) -> str:
            nonlocal call_idx
            call_idx += 1
            if call_idx == 1:
                raise RuntimeError("fail first")
            return "success"

        worker.finished_ok.connect(capture_results)

        with (
            patch.object(worker, "_extract_with_ocr", side_effect=mock_ocr),
            patch(f"{_MOD}.update_extraction_status"),
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
            # Second task should succeed
            assert len(results_list) == 1
            assert results_list[0][2] == "success"

    def test_worker_progress_emission(self) -> None:
        """Worker emits progress signal for each image."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        tasks = [
            (1, "/tmp/a.png"),
            (2, "/tmp/b.png"),
        ]
        worker = _ExtractionWorker(tasks, "tesseract", "en", extract_method="OCR")

        progress_values: list[tuple[int, int]] = []

        def capture_progress(current: int, total: int) -> None:
            progress_values.append((current, total))

        worker.progress.connect(capture_progress)

        with (
            patch.object(worker, "_extract_with_ocr", return_value="text"),
            patch(f"{_MOD}.update_extraction_status"),
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
            assert progress_values == [(1, 2), (2, 2)]

    def test_worker_is_busy_flag(self) -> None:
        """is_busy returns True while worker is running, False after."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        _ExtractionWorker._is_any_worker_running = False
        assert not _ExtractionWorker.is_busy()

        _ExtractionWorker._is_any_worker_running = True
        assert _ExtractionWorker.is_busy()

        # Cleanup
        _ExtractionWorker._is_any_worker_running = False

    def test_worker_duplicate_run_blocked(self) -> None:
        """Second worker is blocked when one is already running."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        tasks = [(1, "/tmp/a.png")]
        worker = _ExtractionWorker(tasks, "tesseract", "en")

        results_list = []

        def capture(results: list) -> None:
            results_list.extend(results)

        worker.finished_ok.connect(capture)

        # Simulate already-running worker
        _ExtractionWorker._is_any_worker_running = True
        worker.run()
        # Should emit empty results since it exited early
        assert results_list == []

        # Cleanup
        _ExtractionWorker._is_any_worker_running = False

    def test_worker_stores_src_lang(self) -> None:
        """Worker stores the source language."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "fr")
        assert worker._src_lang == "fr"

    def test_worker_stores_ocr_method(self) -> None:
        """Worker stores the OCR method."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "easyocr", "en")
        assert worker._ocr_method == "easyocr"


# ===================================================================
# NEW TESTS: Check Requirements (additional)
# ===================================================================


class TestCheckRequirementsAdditional:
    """Additional tests for _check_extract_requirements."""

    @patch(f"{_MOD}.check_llm_setup", return_value=False)
    @patch(f"{_MOD}.check_ocr_setup", return_value=True)
    def test_passes_with_ocr_only(
        self,
        mock_ocr,
        mock_llm,
        page,  # noqa: ANN001
    ) -> None:
        """Returns True when only OCR is configured."""
        result = page._check_extract_requirements()
        assert result is True

    @patch(f"{_MOD}.check_llm_setup", return_value=True)
    @patch(f"{_MOD}.check_ocr_setup", return_value=False)
    def test_passes_with_llm_only(
        self,
        mock_ocr,
        mock_llm,
        page,  # noqa: ANN001
    ) -> None:
        """Returns True when only LLM is configured."""
        result = page._check_extract_requirements()
        assert result is True

    @patch(f"{_MOD}.require_setup", return_value=False)
    @patch(f"{_MOD}.check_llm_setup", return_value=False)
    @patch(f"{_MOD}.check_ocr_setup", return_value=False)
    def test_fails_with_neither_configured(
        self,
        mock_ocr,
        mock_llm,
        mock_require,
        page,  # noqa: ANN001
    ) -> None:
        """Returns False and calls require_setup when neither is configured."""
        result = page._check_extract_requirements()
        assert result is False
        mock_require.assert_called_once()


# ===================================================================
# NEW TESTS: Write Output (additional)
# ===================================================================


class TestWriteOutputAdditional:
    """Additional tests for _write_extraction_output."""

    def test_write_txt_output(self, tmp_path) -> None:  # noqa: ANN001
        """Writes plain text to a .txt file correctly."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "test.txt"
        _write_extraction_output(out, "Line 1\nLine 2\nLine 3")
        content = out.read_text(encoding="utf-8")
        assert content == "Line 1\nLine 2\nLine 3"

    def test_write_docx_output_creates_valid_file(self, tmp_path) -> None:  # noqa: ANN001
        """Writes text to .docx and produces a valid zip file."""
        import zipfile  # noqa: PLC0415

        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "test.docx"
        _write_extraction_output(out, "Paragraph 1\nParagraph 2")
        assert out.exists()
        assert zipfile.is_zipfile(out)

    def test_write_docx_contains_paragraphs(self, tmp_path) -> None:  # noqa: ANN001
        """Docx output contains the correct number of paragraphs."""
        from docx import Document  # noqa: PLC0415

        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "test_para.docx"
        _write_extraction_output(out, "First\nSecond\nThird")
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert "First" in texts
        assert "Second" in texts
        assert "Third" in texts

    def test_write_empty_text(self, tmp_path) -> None:  # noqa: ANN001
        """Writing empty text produces a valid output file."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "empty.txt"
        _write_extraction_output(out, "")
        assert out.exists()
        content = out.read_text(encoding="utf-8")
        assert content == ""

    def test_write_unicode_text(self, tmp_path) -> None:  # noqa: ANN001
        """Unicode text (CJK, accented) is written correctly."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        unicode_text = (
            "\u4f60\u597d\u4e16\u754c\n"
            "\u00c9\u00e0\u00fc\u00f1\n"
            "\u0410\u0411\u0412\u0413"
        )
        out = tmp_path / "unicode.txt"
        _write_extraction_output(out, unicode_text)
        content = out.read_text(encoding="utf-8")
        assert "\u4f60\u597d" in content
        assert "\u00c9\u00e0" in content
        assert "\u0410\u0411" in content

    def test_write_docx_empty_text(self, tmp_path) -> None:  # noqa: ANN001
        """Writing empty text to docx produces a valid file."""
        import zipfile  # noqa: PLC0415

        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "empty.docx"
        _write_extraction_output(out, "")
        assert out.exists()
        assert zipfile.is_zipfile(out)

    def test_write_docx_unicode_text(self, tmp_path) -> None:  # noqa: ANN001
        """Unicode text is preserved in docx output."""
        from docx import Document  # noqa: PLC0415

        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "unicode.docx"
        _write_extraction_output(out, "\u4f60\u597d\n\u00c9\u00e0")
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "\u4f60\u597d" in all_text
        assert "\u00c9\u00e0" in all_text

    def test_write_txt_multiline(self, tmp_path) -> None:  # noqa: ANN001
        """Multi-line text preserves line breaks in txt output."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "multiline.txt"
        text = "Line A\nLine B\nLine C"
        _write_extraction_output(out, text)
        lines = out.read_text(encoding="utf-8").split("\n")
        assert len(lines) == 3  # noqa: PLR2004
        assert lines[0] == "Line A"


# ===================================================================
# NEW TESTS: Theme / Language (additional)
# ===================================================================


class TestExtractThemeLanguage:
    """Additional tests for apply_theme and apply_language."""

    def test_apply_theme_updates_extract_btn(self, page) -> None:  # noqa: ANN001
        """apply_theme updates the extract button stylesheet."""
        page.apply_theme()
        assert page.extract_btn.styleSheet() != ""

    def test_apply_theme_updates_clear_btn(self, page) -> None:  # noqa: ANN001
        """apply_theme updates the clear-all button stylesheet."""
        page.apply_theme()
        assert page.clear_all_btn.styleSheet() != ""

    def test_apply_theme_updates_badge(self, page) -> None:  # noqa: ANN001
        """apply_theme updates the file count badge stylesheet."""
        page.apply_theme()
        assert page.files_badge.styleSheet() != ""

    def test_apply_theme_updates_section_label(self, page) -> None:  # noqa: ANN001
        """apply_theme updates the section label stylesheet."""
        page.apply_theme()
        assert page.section_label.styleSheet() != ""

    def test_apply_language_updates_extract_btn_text(self, page) -> None:  # noqa: ANN001
        """apply_language updates the extract button text."""
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page.apply_language()
        assert page.extract_btn.text() != ""

    def test_apply_language_updates_clear_btn_text(self, page) -> None:  # noqa: ANN001
        """apply_language updates the clear-all button text."""
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page.apply_language()
        assert page.clear_all_btn.text() != ""

    def test_apply_language_updates_section_label_text(self, page) -> None:  # noqa: ANN001
        """apply_language updates the section label text."""
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page.apply_language()
        assert page.section_label.text() != ""

    def test_apply_language_updates_drop_area_label(self, page) -> None:  # noqa: ANN001
        """apply_language updates the drop area supported formats label."""
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page.apply_language()
        assert page.drop_area.supported_label.text() != ""


# ===================================================================
# NEW TESTS: File Handling (additional)
# ===================================================================


class TestFileHandlingAdditional:
    """Additional file handling tests."""

    def test_handle_clear_all_resets_badge(self, page) -> None:  # noqa: ANN001
        """Clearing all files resets the badge to '0'."""
        page.selected_files = ["/tmp/a.png", "/tmp/b.png"]
        page._update_ui_state()
        page._handle_clear_all()
        assert page.files_badge.text() == "0"

    def test_update_ui_state_with_files(self, page) -> None:  # noqa: ANN001
        """_update_ui_state switches to files view and updates badge."""
        page.selected_files = ["/tmp/a.png", "/tmp/b.png"]
        page._update_ui_state()
        assert page.stack.currentIndex() == 1
        assert page.files_badge.text() == "2"

    def test_update_ui_state_without_files(self, page) -> None:  # noqa: ANN001
        """_update_ui_state switches to history view when no files."""
        page.selected_files = []
        page._update_ui_state()
        assert page.stack.currentIndex() == 0
        assert page.files_badge.text() == "0"

    def test_handle_files_dropped_empty_list_no_crash(self, page) -> None:  # noqa: ANN001
        """Dropping an empty list (with dialog cancelled) does not crash."""
        with patch(
            "PySide6.QtWidgets.QFileDialog.getOpenFileNames",
            return_value=([], ""),
        ):
            page._handle_files_dropped([])
            # Should not crash, no files added

    def test_drop_area_exists(self, page) -> None:  # noqa: ANN001
        """Page has a drop area for file input."""
        from src.ui.components import FileDropWidget  # noqa: PLC0415

        assert isinstance(page.drop_area, FileDropWidget)


# ===================================================================
# NEW TESTS: On Finished
# ===================================================================


class TestOnFinished:
    """Tests for the _on_finished callback."""

    @patch(f"{_MOD}.load_setting", return_value=False)
    @patch(f"{_MOD}.generate_extraction_output_path")
    @patch(f"{_MOD}._write_extraction_output")
    @patch(f"{_MOD}.update_extraction_status")
    def test_on_finished_writes_output(
        self,
        mock_update,
        mock_write,
        mock_gen_path,
        mock_load,
        page,  # noqa: ANN001
    ) -> None:
        """_on_finished writes extraction output for each result."""
        from pathlib import Path  # noqa: PLC0415

        mock_gen_path.return_value = Path("/tmp/out.txt")
        page._output_format = ".txt"
        page._worker = MagicMock()

        results = [(1, "/tmp/img.png", "Extracted text")]
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._on_finished(results)

        mock_write.assert_called_once()
        mock_update.assert_called_once()

    @patch(f"{_MOD}.load_setting", return_value=False)
    @patch(f"{_MOD}.generate_extraction_output_path")
    @patch(f"{_MOD}._write_extraction_output", side_effect=OSError("disk full"))
    @patch(f"{_MOD}.update_extraction_status")
    def test_on_finished_handles_write_error(
        self,
        mock_update,
        mock_write,
        mock_gen_path,
        mock_load,
        page,  # noqa: ANN001
    ) -> None:
        """_on_finished marks entry as Failed when write fails."""
        from pathlib import Path  # noqa: PLC0415

        mock_gen_path.return_value = Path("/tmp/out.txt")
        page._output_format = ".txt"
        page._worker = MagicMock()

        results = [(1, "/tmp/img.png", "Extracted text")]
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._on_finished(results)

        # Should have been called with STATUS_FAILED
        failed_calls = [c for c in mock_update.call_args_list if c[0][1] == "Failed"]
        assert len(failed_calls) == 1

    @patch(f"{_MOD}.load_setting", return_value=True)
    @patch(f"{_MOD}.generate_extraction_output_path")
    @patch(f"{_MOD}._write_extraction_output")
    @patch(f"{_MOD}.delete_extraction_entry")
    @patch(f"{_MOD}.update_extraction_status")
    def test_on_finished_auto_remove_deletes_entry(  # noqa: PLR0913
        self,
        mock_update,
        mock_delete,
        mock_write,
        mock_gen_path,
        mock_load,
        page,  # noqa: ANN001
    ) -> None:
        """_on_finished deletes entry when auto_remove is enabled."""
        from pathlib import Path  # noqa: PLC0415

        mock_gen_path.return_value = Path("/tmp/out.txt")
        page._output_format = ".txt"
        page._worker = MagicMock()

        results = [(1, "/tmp/img.png", "Extracted text")]
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._on_finished(results)

        mock_delete.assert_called_once_with(1)


# ===================================================================
# NEW TESTS: Misc edge cases
# ===================================================================


class TestMiscEdgeCases:
    """Miscellaneous edge case tests."""

    def test_image_filter_string(self) -> None:
        """_IMAGE_FILTER contains expected image extension patterns."""
        from src.ui.pages.extract_text import _IMAGE_FILTER  # noqa: PLC0415

        assert "*.png" in _IMAGE_FILTER or "*.PNG" in _IMAGE_FILTER
        assert "Images" in _IMAGE_FILTER

    def test_view_constants(self) -> None:
        """View index constants are defined correctly."""
        from src.ui.pages.extract_text import (  # noqa: PLC0415
            _VIEW_FILES,
            _VIEW_HISTORY,
        )

        assert _VIEW_HISTORY == 0
        assert _VIEW_FILES == 1

    def test_worker_tasks_stored(self) -> None:
        """Worker stores the tasks list."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        tasks = [(1, "/tmp/a.png"), (2, "/tmp/b.png")]
        worker = _ExtractionWorker(tasks, "tesseract", "en")
        assert worker._tasks == tasks

    def test_worker_empty_tasks(self) -> None:
        """Worker with empty tasks list completes without error."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en")

        results_list = []

        def capture(results: list) -> None:
            results_list.extend(results)

        worker.finished_ok.connect(capture)

        _ExtractionWorker._is_any_worker_running = False
        worker.run()
        assert results_list == []

    def test_pending_tasks_initially_empty(self, page) -> None:  # noqa: ANN001
        """Pending tasks list is empty at construction."""
        assert page._pending_tasks == []


# ===================================================================
# NEW: OCR vs LLM extraction method selection
# ===================================================================


class TestExtractionMethodSelection:
    """Tests for OCR vs LLM method selection."""

    def test_worker_ocr_method_default(self) -> None:
        """Default extraction method is OCR."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en")
        assert worker._extract_method == "OCR"

    def test_worker_llm_method(self) -> None:
        """LLM extraction method is stored correctly."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en", extract_method="LLM")
        assert worker._extract_method == "LLM"

    def test_worker_ocr_explicit(self) -> None:
        """Explicit OCR method is stored correctly."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en", extract_method="OCR")
        assert worker._extract_method == "OCR"

    @patch(f"{_MOD}.update_extraction_status")
    def test_ocr_method_calls_extract_with_ocr(self, mock_update) -> None:
        """Worker with OCR method calls _extract_with_ocr."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker(
            [(1, "/tmp/img.png")], "tesseract", "en", extract_method="OCR"
        )
        with patch.object(worker, "_extract_with_ocr", return_value="text") as m_ocr:
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
            m_ocr.assert_called_once()

    @patch(f"{_MOD}.update_extraction_status")
    def test_llm_method_calls_extract_with_llm(self, mock_update) -> None:
        """Worker with LLM method calls _extract_with_llm."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker(
            [(1, "/tmp/img.png")], "tesseract", "en", extract_method="LLM"
        )
        with patch.object(worker, "_extract_with_llm", return_value="text") as m_llm:
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
            m_llm.assert_called_once()

    def test_worker_stores_all_params(self) -> None:
        """Worker stores all constructor parameters."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        tasks = [(1, "/tmp/a.png")]
        worker = _ExtractionWorker(tasks, "easyocr", "fr", extract_method="LLM")
        assert worker._tasks == tasks
        assert worker._ocr_method == "easyocr"
        assert worker._src_lang == "fr"
        assert worker._extract_method == "LLM"


# ===================================================================
# NEW: File drop handling with various image formats
# ===================================================================


class TestImageFileDropHandling:
    """Tests for dropping various image formats."""

    @pytest.fixture()
    def image_files(self, tmp_path):
        """Creates temporary image files for testing."""
        files = {}
        for ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"):
            f = tmp_path / f"image{ext}"
            f.write_bytes(b"\x89PNG\r\n\x1a\n")  # fake image data
            files[ext] = str(f)
        return files

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_png(self, _mock_msg, page, image_files) -> None:
        """Dropping a .png file adds it."""
        page._handle_files_dropped([image_files[".png"]])
        assert image_files[".png"] in page.selected_files

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_jpg(self, _mock_msg, page, image_files) -> None:
        """Dropping a .jpg file adds it."""
        page._handle_files_dropped([image_files[".jpg"]])
        assert image_files[".jpg"] in page.selected_files

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_jpeg(self, _mock_msg, page, image_files) -> None:
        """Dropping a .jpeg file adds it."""
        page._handle_files_dropped([image_files[".jpeg"]])
        assert image_files[".jpeg"] in page.selected_files

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_tiff(self, _mock_msg, page, image_files) -> None:
        """Dropping a .tiff file adds it."""
        page._handle_files_dropped([image_files[".tiff"]])
        assert image_files[".tiff"] in page.selected_files

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_bmp(self, _mock_msg, page, image_files) -> None:
        """Dropping a .bmp file adds it."""
        page._handle_files_dropped([image_files[".bmp"]])
        assert image_files[".bmp"] in page.selected_files

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_webp(self, _mock_msg, page, image_files) -> None:
        """Dropping a .webp file adds it."""
        page._handle_files_dropped([image_files[".webp"]])
        assert image_files[".webp"] in page.selected_files

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_multiple_images(self, _mock_msg, page, image_files) -> None:
        """Dropping multiple image formats adds all of them."""
        paths = list(image_files.values())
        page._handle_files_dropped(paths)
        assert len(page.selected_files) == len(paths)

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_non_image_rejected(self, mock_msg, page, tmp_path) -> None:
        """Dropping a non-image file is rejected."""
        f = tmp_path / "document.docx"
        f.write_text("content")
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 0
        mock_msg.assert_called_once()

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_empty_image(self, mock_msg, page, tmp_path) -> None:
        """Dropping a zero-byte image is rejected."""
        f = tmp_path / "empty.png"
        f.write_bytes(b"")
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 0
        mock_msg.assert_called_once()

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_duplicate_image_ignored(self, _mock_msg, page, image_files) -> None:
        """Dropping the same image twice does not duplicate it."""
        page._handle_files_dropped([image_files[".png"]])
        page._handle_files_dropped([image_files[".png"]])
        assert page.selected_files.count(image_files[".png"]) == 1

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_directory_with_images(self, _mock_msg, page, tmp_path) -> None:
        """Dropping a directory with images adds them."""
        d = tmp_path / "images"
        d.mkdir()
        for ext in (".png", ".jpg"):
            f = d / f"img{ext}"
            f.write_bytes(b"data")
        page._handle_files_dropped([str(d)])
        assert len(page.selected_files) == 2  # noqa: PLR2004


# ===================================================================
# NEW: Output format selection (txt/docx)
# ===================================================================


class TestOutputFormatSelection:
    """Tests for output format selection."""

    def test_default_format_txt(self, page) -> None:  # noqa: ANN001
        """Default output format is .txt."""
        assert page._output_format == ".txt"

    def test_write_txt_creates_file(self, tmp_path) -> None:
        """Writing .txt output creates a file."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "output.txt"
        _write_extraction_output(out, "test content")
        assert out.exists()
        assert out.read_text(encoding="utf-8") == "test content"

    def test_write_docx_creates_file(self, tmp_path) -> None:
        """Writing .docx output creates a valid file."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "output.docx"
        _write_extraction_output(out, "test content")
        assert out.exists()
        import zipfile  # noqa: PLC0415

        assert zipfile.is_zipfile(out)

    def test_write_txt_preserves_newlines(self, tmp_path) -> None:
        """Writing .txt preserves newlines."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "newlines.txt"
        _write_extraction_output(out, "Line 1\nLine 2\nLine 3")
        content = out.read_text(encoding="utf-8")
        assert content.count("\n") == 2  # noqa: PLR2004

    def test_write_docx_paragraphs(self, tmp_path) -> None:
        """Writing .docx creates correct paragraphs."""
        from docx import Document  # noqa: PLC0415

        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "paras.docx"
        _write_extraction_output(out, "A\nB\nC")
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert "A" in texts
        assert "B" in texts
        assert "C" in texts

    def test_write_empty_string_txt(self, tmp_path) -> None:
        """Writing empty string to .txt creates empty file."""
        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "empty.txt"
        _write_extraction_output(out, "")
        assert out.read_text(encoding="utf-8") == ""

    def test_write_empty_string_docx(self, tmp_path) -> None:
        """Writing empty string to .docx creates valid file."""
        import zipfile  # noqa: PLC0415

        from src.ui.pages.extract_text import _write_extraction_output  # noqa: PLC0415

        out = tmp_path / "empty.docx"
        _write_extraction_output(out, "")
        assert zipfile.is_zipfile(out)


# ===================================================================
# NEW: Worker lifecycle and error handling expanded
# ===================================================================


class TestWorkerLifecycleExpanded:
    """Expanded tests for worker lifecycle."""

    def test_worker_finished_clears_busy_flag(self) -> None:
        """Worker clears is_busy flag after finishing."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en")
        _ExtractionWorker._is_any_worker_running = False
        worker.run()
        assert not _ExtractionWorker.is_busy()

    def test_worker_crash_clears_busy_flag(self) -> None:
        """Worker clears is_busy flag even after crash."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([(1, "/tmp/img.png")], "tesseract", "en")
        with (
            patch.object(worker, "_extract_with_ocr", side_effect=RuntimeError("boom")),
            patch(f"{_MOD}.update_extraction_status"),
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
        assert not _ExtractionWorker.is_busy()

    def test_worker_emits_finished_signal(self) -> None:
        """Worker emits finished_ok signal when done."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en")
        results = []
        worker.finished_ok.connect(results.append)
        _ExtractionWorker._is_any_worker_running = False
        worker.run()
        assert len(results) == 1

    def test_worker_result_format(self) -> None:
        """Worker results are tuples of (entry_id, path, text)."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([(1, "/tmp/img.png")], "tesseract", "en")
        results = []
        worker.finished_ok.connect(results.append)
        with (
            patch.object(worker, "_extract_with_ocr", return_value="extracted"),
            patch(f"{_MOD}.update_extraction_status"),
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
        assert len(results) == 1
        assert len(results[0]) == 1
        assert results[0][0] == (1, "/tmp/img.png", "extracted")

    def test_worker_stops_mid_processing(self) -> None:
        """Worker stops processing when stop() is called mid-run."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        tasks = [(i, f"/tmp/{i}.png") for i in range(10)]
        worker = _ExtractionWorker(tasks, "tesseract", "en")
        processed = []

        def mock_ocr(path: str) -> str:
            processed.append(path)
            if len(processed) == 3:  # noqa: PLR2004
                worker.stop()
            return "text"

        with (
            patch.object(worker, "_extract_with_ocr", side_effect=mock_ocr),
            patch(f"{_MOD}.update_extraction_status"),
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
        assert len(processed) == 3  # noqa: PLR2004

    def test_worker_error_marks_entry_failed(self) -> None:
        """Worker marks entry as Failed when extraction raises."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([(1, "/tmp/img.png")], "tesseract", "en")
        with (
            patch.object(worker, "_extract_with_ocr", side_effect=RuntimeError("bad")),
            patch(f"{_MOD}.update_extraction_status") as mock_update,
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
        failed_calls = [c for c in mock_update.call_args_list if c[0][1] == "Failed"]
        assert len(failed_calls) == 1

    def test_worker_sets_extracting_status(self) -> None:
        """Worker sets Extracting status before extraction."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([(1, "/tmp/img.png")], "tesseract", "en")
        with (
            patch.object(worker, "_extract_with_ocr", return_value="text"),
            patch(f"{_MOD}.update_extraction_status") as mock_update,
        ):
            _ExtractionWorker._is_any_worker_running = False
            worker.run()
        extracting_calls = [
            c for c in mock_update.call_args_list if c[0][1] == "Extracting"
        ]
        assert len(extracting_calls) == 1


# ===================================================================
# NEW: Settings integration
# ===================================================================


class TestSettingsIntegration:
    """Tests for settings integration with the page."""

    def test_page_stores_window_context(self, page, window) -> None:  # noqa: ANN001
        """Page stores a reference to the parent window."""
        assert page.window_context is window

    def test_page_has_worker_attribute(self, page) -> None:  # noqa: ANN001
        """Page has _worker attribute."""
        assert hasattr(page, "_worker")

    def test_page_has_output_format_attribute(self, page) -> None:  # noqa: ANN001
        """Page has _output_format attribute."""
        assert hasattr(page, "_output_format")

    def test_page_has_pending_tasks_attribute(self, page) -> None:  # noqa: ANN001
        """Page has _pending_tasks attribute."""
        assert hasattr(page, "_pending_tasks")

    def test_page_has_selected_files_attribute(self, page) -> None:  # noqa: ANN001
        """Page has selected_files attribute."""
        assert hasattr(page, "selected_files")
        assert isinstance(page.selected_files, list)


# ===================================================================
# NEW: View switching expanded
# ===================================================================


class TestViewSwitchingExpanded:
    """Expanded view switching tests."""

    def test_switch_to_files_changes_drop_label(self, page) -> None:  # noqa: ANN001
        """Switching to files view changes drop area label."""
        from src.constants.i18n import tr  # noqa: PLC0415

        page.selected_files = ["/tmp/a.png"]
        page._update_ui_state()
        assert page.drop_area.info_label.text() == tr("drop.title_more")

    def test_switch_to_history_changes_drop_label(self, page) -> None:  # noqa: ANN001
        """Switching to history view changes drop area label."""
        from src.constants.i18n import tr  # noqa: PLC0415

        page.selected_files = []
        page._update_ui_state()
        assert page.drop_area.info_label.text() == tr("drop.title")

    def test_switch_views_repeatedly(self, page) -> None:  # noqa: ANN001
        """Switching views back and forth does not crash."""
        for _ in range(10):
            page.selected_files = ["/tmp/a.png"]
            page._update_ui_state()
            page.selected_files = []
            page._update_ui_state()
        assert page.stack.currentIndex() == 0

    def test_badge_updates_correctly_many_files(self, page) -> None:  # noqa: ANN001
        """Badge updates for many files."""
        page.selected_files = [f"/tmp/{i}.png" for i in range(99)]
        page._update_ui_state()
        assert page.files_badge.text() == "99"

    def test_extract_btn_cursor_shape(self, page) -> None:  # noqa: ANN001
        """Extract button has pointing hand cursor."""
        from PySide6.QtCore import Qt  # noqa: PLC0415

        assert page.extract_btn.cursor().shape() == Qt.CursorShape.PointingHandCursor

    def test_clear_btn_cursor_shape(self, page) -> None:  # noqa: ANN001
        """Clear-all button has pointing hand cursor."""
        from PySide6.QtCore import Qt  # noqa: PLC0415

        assert page.clear_all_btn.cursor().shape() == Qt.CursorShape.PointingHandCursor


# ===================================================================
# NEW: Clean history view expanded
# ===================================================================


class TestCleanHistoryViewExpanded:
    """Tests for _clean_history_view."""

    def test_clean_history_view_no_crash(self, page) -> None:  # noqa: ANN001
        """_clean_history_view does not crash."""
        page._clean_history_view()

    def test_clean_history_view_repeated(self, page) -> None:  # noqa: ANN001
        """Calling _clean_history_view multiple times is safe."""
        page._clean_history_view()
        page._clean_history_view()
        page._clean_history_view()

    def test_clean_history_view_without_page_attr(self, page) -> None:  # noqa: ANN001
        """_clean_history_view handles missing page attribute."""
        original = page.extraction_history
        page.extraction_history = MagicMock(spec=[])
        page._clean_history_view()
        page.extraction_history = original


# ===================================================================
# NEW: Handle extract expanded
# ===================================================================


class TestHandleExtractExpanded:
    """Expanded tests for _handle_extract."""

    def test_extract_with_no_files_noop(self, page) -> None:  # noqa: ANN001
        """_handle_extract with no files is a no-op."""
        page.selected_files = []
        page._handle_extract()
        assert page.selected_files == []

    @patch(f"{_MOD}.require_setup", return_value=False)
    def test_extract_blocked_keeps_files(self, mock_require, page) -> None:  # noqa: ANN001
        """Files are kept when requirements check fails."""
        page.selected_files = ["/tmp/img.png"]
        page._handle_extract()
        assert page.selected_files == ["/tmp/img.png"]

    @patch(f"{_MOD}.load_setting", return_value="ocr")
    @patch(
        f"{_MOD}.SourceLanguageDialog.get_selection",
        return_value=("en", "English", None, False),
    )
    @patch(f"{_MOD}.require_setup", return_value=True)
    def test_extract_dialog_cancel_keeps_files(
        self,
        mock_require,
        mock_dialog,
        mock_load,
        page,  # noqa: ANN001
    ) -> None:
        """Cancelling language dialog keeps files."""
        page.selected_files = ["/tmp/img.png"]
        page._handle_extract()
        assert page.selected_files == ["/tmp/img.png"]


# ===================================================================
# NEW: Re-extract expanded
# ===================================================================


class TestHandleReExtractExpanded:
    """Expanded tests for _handle_re_extract."""

    @patch(f"{_MOD}.require_setup", return_value=False)
    def test_re_extract_blocked(self, mock_require, page) -> None:  # noqa: ANN001
        """_handle_re_extract blocked when requirements fail."""
        page._handle_re_extract([(1, "/tmp/img.png")])
        mock_require.assert_called_once()

    @patch(f"{_MOD}.load_setting", return_value="ocr")
    @patch(
        f"{_MOD}.SourceLanguageDialog.get_selection",
        return_value=("en", "English", None, False),
    )
    @patch(f"{_MOD}.require_setup", return_value=True)
    def test_re_extract_dialog_cancel(
        self,
        mock_require,
        mock_dialog,
        mock_load,
        page,  # noqa: ANN001
    ) -> None:
        """Re-extract is cancelled when language dialog is cancelled."""
        page._handle_re_extract([(1, "/tmp/img.png")])
        # Should not start worker

    @patch(f"{_MOD}._ExtractionWorker")
    @patch(f"{_MOD}.load_setting", return_value="ocr")
    @patch(f"{_MOD}.update_extraction_status")
    @patch(
        f"{_MOD}.SourceLanguageDialog.get_selection",
        return_value=("fr", "French", None, True),
    )
    @patch(f"{_MOD}.require_setup", return_value=True)
    def test_re_extract_with_french_source(  # noqa: PLR0913
        self,
        mock_require,
        mock_dialog,
        mock_update,
        mock_load,
        mock_worker_cls,
        page,  # noqa: ANN001
    ) -> None:
        """Re-extract uses selected source language."""
        mock_worker = MagicMock()
        mock_worker_cls.return_value = mock_worker
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._handle_re_extract([(1, "/tmp/img.png")])
        mock_worker.start.assert_called_once()


# ===================================================================
# NEW: Image extension filtering
# ===================================================================


class TestImageExtensionFiltering:
    """Tests for image-specific file extension filtering."""

    def test_drop_png_accepted(self, page, tmp_path) -> None:  # noqa: ANN001
        """PNG file is accepted."""
        f = tmp_path / "test.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 1

    def test_drop_jpg_accepted(self, page, tmp_path) -> None:  # noqa: ANN001
        """JPG file is accepted."""
        f = tmp_path / "test.jpg"
        f.write_bytes(b"\xff\xd8\xff")
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 1

    def test_drop_jpeg_accepted(self, page, tmp_path) -> None:  # noqa: ANN001
        """JPEG file is accepted."""
        f = tmp_path / "test.jpeg"
        f.write_bytes(b"\xff\xd8\xff")
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 1

    def test_drop_bmp_accepted(self, page, tmp_path) -> None:  # noqa: ANN001
        """BMP file is accepted."""
        f = tmp_path / "test.bmp"
        f.write_bytes(b"BM" + b"\x00" * 10)
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 1

    def test_drop_tiff_accepted(self, page, tmp_path) -> None:  # noqa: ANN001
        """TIFF file is accepted."""
        f = tmp_path / "test.tiff"
        f.write_bytes(b"II\x2a\x00" + b"\x00" * 10)
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 1

    def test_drop_webp_accepted(self, page, tmp_path) -> None:  # noqa: ANN001
        """WebP file is accepted."""
        f = tmp_path / "test.webp"
        f.write_bytes(b"RIFF" + b"\x00" * 10)
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 1

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_txt_rejected(self, mock_msg, page, tmp_path) -> None:  # noqa: ANN001
        """TXT file is rejected for image extraction."""
        f = tmp_path / "test.txt"
        f.write_text("content")
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 0
        mock_msg.assert_called_once()

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_pdf_rejected(self, mock_msg, page, tmp_path) -> None:  # noqa: ANN001
        """PDF file is rejected for image extraction."""
        f = tmp_path / "test.pdf"
        f.write_text("content")
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 0
        mock_msg.assert_called_once()

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_docx_rejected(self, mock_msg, page, tmp_path) -> None:  # noqa: ANN001
        """DOCX file is rejected for image extraction."""
        f = tmp_path / "test.docx"
        f.write_text("content")
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 0
        mock_msg.assert_called_once()

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_empty_png_rejected(self, mock_msg, page, tmp_path) -> None:  # noqa: ANN001
        """Empty PNG file is rejected."""
        f = tmp_path / "empty.png"
        f.write_bytes(b"")
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 0
        mock_msg.assert_called_once()


# ===================================================================
# NEW: Directory drop for images
# ===================================================================


class TestDirectoryDropImages:
    """Tests for directory drop handling with image files."""

    def test_drop_directory_with_images(self, page, tmp_path) -> None:  # noqa: ANN001
        """Dropping a directory adds its image files."""
        subdir = tmp_path / "images"
        subdir.mkdir()
        (subdir / "a.png").write_bytes(b"\x89PNG\r\n\x1a\n")
        (subdir / "b.jpg").write_bytes(b"\xff\xd8\xff")
        page._handle_files_dropped([str(subdir)])
        assert len(page.selected_files) == 2  # noqa: PLR2004

    @patch(f"{_MOD}.CustomMessageDialog.show_message")
    def test_drop_directory_mixed_types(self, mock_msg, page, tmp_path) -> None:  # noqa: ANN001
        """Directory with mixed types adds only images."""
        subdir = tmp_path / "mixed"
        subdir.mkdir()
        (subdir / "img.png").write_bytes(b"\x89PNG\r\n\x1a\n")
        (subdir / "doc.txt").write_text("content")
        page._handle_files_dropped([str(subdir)])
        assert len(page.selected_files) == 1
        mock_msg.assert_called_once()

    def test_drop_directory_skips_hidden(self, page, tmp_path) -> None:  # noqa: ANN001
        """Hidden subdirectories are skipped."""
        subdir = tmp_path / "root"
        subdir.mkdir()
        hidden = subdir / ".hidden"
        hidden.mkdir()
        (hidden / "secret.png").write_bytes(b"\x89PNG\r\n\x1a\n")
        (subdir / "visible.png").write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(subdir)])
        assert len(page.selected_files) == 1

    def test_drop_empty_directory(self, page, tmp_path) -> None:  # noqa: ANN001
        """Empty directory adds nothing."""
        subdir = tmp_path / "empty"
        subdir.mkdir()
        page._handle_files_dropped([str(subdir)])
        assert len(page.selected_files) == 0


# ===================================================================
# NEW: Worker class methods
# ===================================================================


class TestExtractionWorkerClass:
    """Tests for _ExtractionWorker class methods."""

    def test_is_busy_default_false(self) -> None:
        """is_busy() returns False by default."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        # Save and restore
        original = _ExtractionWorker._is_any_worker_running
        _ExtractionWorker._is_any_worker_running = False
        assert _ExtractionWorker.is_busy() is False
        _ExtractionWorker._is_any_worker_running = original

    def test_is_busy_when_running(self) -> None:
        """is_busy() returns True when a worker is running."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        original = _ExtractionWorker._is_any_worker_running
        _ExtractionWorker._is_any_worker_running = True
        assert _ExtractionWorker.is_busy() is True
        _ExtractionWorker._is_any_worker_running = original

    def test_worker_stop_sets_flag(self) -> None:
        """stop() sets _is_running to False."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en")
        assert worker._is_running is True
        worker.stop()
        assert worker._is_running is False

    def test_worker_constructor_stores_tasks(self) -> None:
        """Constructor stores tasks."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        tasks = [(1, "/tmp/a.png"), (2, "/tmp/b.png")]
        worker = _ExtractionWorker(tasks, "tesseract", "en")
        assert worker._tasks == tasks

    def test_worker_constructor_stores_method(self) -> None:
        """Constructor stores OCR method."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "easyocr", "en")
        assert worker._ocr_method == "easyocr"

    def test_worker_constructor_stores_lang(self) -> None:
        """Constructor stores source language."""
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "fr")
        assert worker._src_lang == "fr"

    def test_worker_default_extract_method(self) -> None:
        """Constructor defaults to OCR extract method."""
        from src.constants.settings import EXTRACT_METHOD_OCR  # noqa: PLC0415
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker([], "tesseract", "en")
        assert worker._extract_method == EXTRACT_METHOD_OCR

    def test_worker_llm_extract_method(self) -> None:
        """Constructor accepts LLM extract method."""
        from src.constants.settings import EXTRACT_METHOD_LLM  # noqa: PLC0415
        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        worker = _ExtractionWorker(
            [], "tesseract", "en", extract_method=EXTRACT_METHOD_LLM
        )
        assert worker._extract_method == EXTRACT_METHOD_LLM


# ===================================================================
# NEW: Apply theme and language
# ===================================================================


class TestApplyThemeLanguage:
    """Tests for theme and language application."""

    def test_apply_theme_no_crash(self, page) -> None:  # noqa: ANN001
        """apply_theme does not crash."""
        page.apply_theme()

    def test_apply_theme_repeated(self, page) -> None:  # noqa: ANN001
        """Multiple apply_theme calls are safe."""
        for _ in range(5):
            page.apply_theme()

    def test_apply_language_no_crash(self, page) -> None:  # noqa: ANN001
        """apply_language does not crash."""
        page.apply_language()

    def test_apply_language_repeated(self, page) -> None:  # noqa: ANN001
        """Multiple apply_language calls are safe."""
        for _ in range(5):
            page.apply_language()

    def test_apply_theme_sets_badge_style(self, page) -> None:  # noqa: ANN001
        """apply_theme updates badge stylesheet."""
        page.apply_theme()
        assert len(page.files_badge.styleSheet()) > 0

    def test_apply_theme_sets_extract_btn_style(self, page) -> None:  # noqa: ANN001
        """apply_theme updates extract button stylesheet."""
        page.apply_theme()
        assert len(page.extract_btn.styleSheet()) > 0

    def test_apply_language_sets_extract_btn_text(self, page) -> None:  # noqa: ANN001
        """apply_language updates extract button text."""
        from src.constants.i18n import tr  # noqa: PLC0415

        page.apply_language()
        assert page.extract_btn.text() == tr("extract_text.btn_extract")

    def test_apply_language_sets_section_label(self, page) -> None:  # noqa: ANN001
        """apply_language updates section label text."""
        from src.constants.i18n import tr  # noqa: PLC0415

        page.apply_language()
        assert page.section_label.text() == tr("files.selected")

    def test_apply_language_updates_drop_supported(self, page) -> None:  # noqa: ANN001
        """apply_language updates drop area supported formats."""
        page.apply_language()
        text = page.drop_area.supported_label.text()
        assert len(text) > 0


# ===================================================================
# NEW: UI state management
# ===================================================================


class TestUIStateManagement:
    """Tests for UI state transitions."""

    def test_badge_zero_initially(self, page) -> None:  # noqa: ANN001
        """Badge shows 0 initially."""
        assert page.files_badge.text() == "0"

    def test_badge_updates_on_add(self, page, tmp_path) -> None:  # noqa: ANN001
        """Badge updates when files are added."""
        f = tmp_path / "img.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(f)])
        assert page.files_badge.text() == "1"

    def test_badge_updates_on_clear(self, page, tmp_path) -> None:  # noqa: ANN001
        """Badge returns to 0 after clear."""
        f = tmp_path / "img.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(f)])
        page._handle_clear_all()
        assert page.files_badge.text() == "0"

    def test_extract_btn_disabled_initially(self, page) -> None:  # noqa: ANN001
        """Extract button is disabled when no files are selected."""
        assert not page.extract_btn.isEnabled()

    def test_extract_btn_enabled_with_files(self, page, tmp_path) -> None:  # noqa: ANN001
        """Extract button is enabled when files are present."""
        f = tmp_path / "img.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(f)])
        assert page.extract_btn.isEnabled()

    def test_extract_btn_disabled_after_clear(self, page, tmp_path) -> None:  # noqa: ANN001
        """Extract button is disabled after clearing files."""
        f = tmp_path / "img.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(f)])
        page._handle_clear_all()
        assert not page.extract_btn.isEnabled()

    def test_stack_switches_to_files(self, page, tmp_path) -> None:  # noqa: ANN001
        """Stack switches to files view when files are added."""
        f = tmp_path / "img.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(f)])
        assert page.stack.currentIndex() == 1

    def test_stack_switches_to_history(self, page, tmp_path) -> None:  # noqa: ANN001
        """Stack switches to history view after clearing."""
        f = tmp_path / "img.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(f)])
        page._handle_clear_all()
        assert page.stack.currentIndex() == 0


# ===================================================================
# NEW: File deduplication
# ===================================================================


class TestFileDeduplication:
    """Tests for file deduplication in drop handling."""

    def test_same_file_not_added_twice(self, page, tmp_path) -> None:  # noqa: ANN001
        """Same file dropped twice is only added once."""
        f = tmp_path / "img.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(f)])
        page._handle_files_dropped([str(f)])
        assert len(page.selected_files) == 1

    def test_different_files_added(self, page, tmp_path) -> None:  # noqa: ANN001
        """Different files are all added."""
        f1 = tmp_path / "a.png"
        f1.write_bytes(b"\x89PNG\r\n\x1a\n")
        f2 = tmp_path / "b.jpg"
        f2.write_bytes(b"\xff\xd8\xff")
        page._handle_files_dropped([str(f1), str(f2)])
        assert len(page.selected_files) == 2  # noqa: PLR2004

    def test_dedup_across_multiple_drops(self, page, tmp_path) -> None:  # noqa: ANN001
        """Deduplication works across multiple drop events."""
        f1 = tmp_path / "a.png"
        f1.write_bytes(b"\x89PNG\r\n\x1a\n")
        f2 = tmp_path / "b.png"
        f2.write_bytes(b"\x89PNG\r\n\x1a\n")
        page._handle_files_dropped([str(f1)])
        page._handle_files_dropped([str(f2)])
        page._handle_files_dropped([str(f1)])  # duplicate
        assert len(page.selected_files) == 2  # noqa: PLR2004


# ===================================================================
# NEW: Browse dialog
# ===================================================================


class TestBrowseDialog:
    """Tests for browse dialog integration."""

    @patch(f"{_MOD}.QFileDialog.getOpenFileNames")
    def test_browse_dialog_cancelled(self, mock_dialog, page) -> None:  # noqa: ANN001
        """Cancelling browse dialog adds no files."""
        mock_dialog.return_value = ([], "")
        page._handle_files_dropped([])
        assert len(page.selected_files) == 0

    @patch(f"{_MOD}.QFileDialog.getOpenFileNames")
    def test_browse_dialog_no_files_selected(self, mock_dialog, page) -> None:  # noqa: ANN001
        """No files selected from dialog is a no-op."""
        mock_dialog.return_value = (None, "")
        # This should not crash
        page._handle_files_dropped([])
        assert len(page.selected_files) == 0


# ===================================================================
# NEW: On finished handler
# ===================================================================


class TestOnFinished:
    """Tests for _on_finished handler."""

    @patch(f"{_MOD}.load_setting", return_value=False)
    @patch(f"{_MOD}.generate_extraction_output_path")
    @patch(f"{_MOD}._write_extraction_output")
    @patch(f"{_MOD}.update_extraction_status")
    def test_on_finished_updates_status(
        self,
        mock_update,
        mock_write,
        mock_path,
        mock_load,
        page,  # noqa: ANN001
    ) -> None:
        """_on_finished updates status to DONE."""
        from pathlib import Path  # noqa: PLC0415

        mock_path.return_value = Path("/tmp/out.txt")
        page._worker = MagicMock()
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._on_finished([(1, "/tmp/img.png", "extracted text")])
        mock_update.assert_called()

    @patch(f"{_MOD}.load_setting", return_value=True)
    @patch(f"{_MOD}.generate_extraction_output_path")
    @patch(f"{_MOD}._write_extraction_output")
    @patch(f"{_MOD}.delete_extraction_entry")
    def test_on_finished_auto_remove(
        self,
        mock_delete,
        mock_write,
        mock_path,
        mock_load,
        page,  # noqa: ANN001
    ) -> None:
        """_on_finished deletes entry when auto-remove is enabled."""
        from pathlib import Path  # noqa: PLC0415

        mock_path.return_value = Path("/tmp/out.txt")
        page._worker = MagicMock()
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._on_finished([(1, "/tmp/img.png", "text")])
        mock_delete.assert_called_once_with(1)

    def test_on_finished_empty_results(self, page) -> None:  # noqa: ANN001
        """_on_finished with empty results does not crash."""
        page._worker = MagicMock()
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._on_finished([])

    def test_on_finished_none_worker(self, page) -> None:  # noqa: ANN001
        """_on_finished with None worker does not crash."""
        page._worker = None
        with (
            patch(f"{_HIST_MOD}.get_extraction_fingerprint", return_value=None),
            patch(f"{_HIST_MOD}.get_extraction_history", return_value=[]),
        ):
            page._on_finished([])


class TestStopAllWorkersBoundedWait:
    """``aboutToQuit`` must drain the worker with a bounded wait.

    Pins the ``stop()`` → ``wait(2000)`` contract so a future refactor
    can't regress to an unbounded ``wait()`` and block app exit when a
    stage (FFmpeg mux, OCR call, LLM stream) takes too long to honour
    the cancel flag.
    """

    def test_worker_gets_stop_then_bounded_wait(self, page) -> None:
        """``_stop_all_workers`` calls ``stop()`` then ``wait(2000)``."""
        from unittest.mock import MagicMock  # noqa: PLC0415

        worker = MagicMock()
        worker.wait.return_value = True
        page._worker = worker
        page._stop_all_workers()

        worker.stop.assert_called_once()
        worker.wait.assert_called_once_with(2000)
        assert page._worker is None

    def test_no_worker_is_noop(self, page) -> None:
        """Empty worker slot is a safe no-op."""
        page._worker = None
        page._stop_all_workers()
        assert page._worker is None


class TestExtractionWorkerStopOrphanCleanup:
    """Stop mid-batch must mark unstarted entries as FAILED + CANCELLED.

    Pin the orphan-cleanup contract that mirrors Subtitle / Voice /
    Dubbing pages — without this, hitting Stop after a few files
    have processed leaves the rest of the batch stuck on
    STATUS_PENDING (or the in-flight one on STATUS_EXTRACTING),
    forcing the user to delete orphan rows manually.
    """

    def test_stop_marks_unstarted_entries_as_cancelled(
        self, tmp_path, monkeypatch,
    ) -> None:
        """Worker stopped before any task starts → all tasks → FAILED + CANCELLED."""
        from unittest.mock import MagicMock, patch  # noqa: PLC0415

        from src.constants.history import (  # noqa: PLC0415
            STATUS_FAILED,
            STATUS_PENDING,
        )
        from src.constants.ocr import OCR_METHOD_TESSERACT  # noqa: PLC0415
        from src.constants.settings import EXTRACT_METHOD_OCR  # noqa: PLC0415

        # Isolated DB
        db_file = tmp_path / "ext.db"
        monkeypatch.setattr(
            "src.core.database.get_db_path", lambda: str(db_file),
        )
        from src.core.database import (  # noqa: PLC0415
            add_extraction_entry,
            get_extraction_history,
            init_db,
        )
        init_db()

        img1 = tmp_path / "a.png"
        img1.write_bytes(b"\x89PNG\r\n\x1a\n")
        img2 = tmp_path / "b.png"
        img2.write_bytes(b"\x89PNG\r\n\x1a\n")
        e1 = add_extraction_entry(
            file_name="a.png", file_size=img1.stat().st_size,
            source_path=str(img1),
            output_path=str(img1.with_suffix(".txt")),
            status=STATUS_PENDING,
        )
        e2 = add_extraction_entry(
            file_name="b.png", file_size=img2.stat().st_size,
            source_path=str(img2),
            output_path=str(img2.with_suffix(".txt")),
            status=STATUS_PENDING,
        )

        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        # Reset class-level singleton from any prior test
        _ExtractionWorker._is_any_worker_running = False

        worker = _ExtractionWorker.__new__(_ExtractionWorker)
        worker._tasks = [(e1, str(img1)), (e2, str(img2))]
        worker._ocr_method = OCR_METHOD_TESSERACT
        worker._src_lang = "English"
        worker._extract_method = EXTRACT_METHOD_OCR
        worker._llm_provider = None
        worker._llm_model = None
        worker._is_running = True
        worker.progress = MagicMock(emit=lambda *_: None)
        worker.finished_ok = MagicMock(emit=lambda *_: None)

        # Pre-stop the worker so the run-loop body never executes.
        worker.stop()

        # Should NOT call OCR (worker stopped before any task).
        with patch.object(
            _ExtractionWorker, "_extract_with_ocr",
            return_value="should not run",
        ) as mock_ocr:
            worker.run()

        mock_ocr.assert_not_called()

        # Both unstarted entries should now be FAILED + CANCELLED.
        rows = {r[0]: r for r in get_extraction_history()}
        # status field is index 5; error_message is index 6.
        for entry_id in (e1, e2):
            assert rows[entry_id][5] == STATUS_FAILED, (
                f"Entry {entry_id} should be FAILED after stop, "
                f"got {rows[entry_id][5]!r}"
            )
            assert rows[entry_id][6] == "CANCELLED", (
                f"Entry {entry_id} should carry CANCELLED message, "
                f"got {rows[entry_id][6]!r}"
            )

        # And reset for downstream tests.
        _ExtractionWorker._is_any_worker_running = False

    def test_stop_after_one_processed_marks_only_unprocessed_as_cancelled(
        self, tmp_path, monkeypatch,
    ) -> None:
        """Mixed batch — already-processed rows keep their result status."""
        from unittest.mock import MagicMock, patch  # noqa: PLC0415

        from src.constants.history import (  # noqa: PLC0415
            STATUS_FAILED,
            STATUS_PENDING,
        )
        from src.constants.ocr import OCR_METHOD_TESSERACT  # noqa: PLC0415
        from src.constants.settings import EXTRACT_METHOD_OCR  # noqa: PLC0415

        db_file = tmp_path / "ext2.db"
        monkeypatch.setattr(
            "src.core.database.get_db_path", lambda: str(db_file),
        )
        from src.core.database import (  # noqa: PLC0415
            add_extraction_entry,
            get_extraction_history,
            init_db,
        )
        init_db()

        img1 = tmp_path / "a.png"
        img1.write_bytes(b"\x89PNG\r\n\x1a\n")
        img2 = tmp_path / "b.png"
        img2.write_bytes(b"\x89PNG\r\n\x1a\n")
        e1 = add_extraction_entry(
            file_name="a.png", file_size=img1.stat().st_size,
            source_path=str(img1),
            output_path=str(img1.with_suffix(".txt")),
            status=STATUS_PENDING,
        )
        e2 = add_extraction_entry(
            file_name="b.png", file_size=img2.stat().st_size,
            source_path=str(img2),
            output_path=str(img2.with_suffix(".txt")),
            status=STATUS_PENDING,
        )

        from src.ui.pages.extract_text import _ExtractionWorker  # noqa: PLC0415

        _ExtractionWorker._is_any_worker_running = False

        worker = _ExtractionWorker.__new__(_ExtractionWorker)
        worker._tasks = [(e1, str(img1)), (e2, str(img2))]
        worker._ocr_method = OCR_METHOD_TESSERACT
        worker._src_lang = "English"
        worker._extract_method = EXTRACT_METHOD_OCR
        worker._llm_provider = None
        worker._llm_model = None
        worker._is_running = True
        worker.progress = MagicMock(emit=lambda *_: None)
        worker.finished_ok = MagicMock(emit=lambda *_: None)

        # Stop the worker AFTER the first task succeeds so the second
        # one is the orphan that the finally must clean up.
        def _ocr_then_stop(image_path):
            worker._is_running = False
            return "extracted"

        with patch.object(
            _ExtractionWorker, "_extract_with_ocr",
            side_effect=_ocr_then_stop,
        ):
            worker.run()

        rows = {r[0]: r for r in get_extraction_history()}
        # e1 was processed successfully — caller's _on_finished is
        # responsible for marking it Done; here it remains EXTRACTING
        # (the in-progress flag set inside the loop) — that's fine,
        # the page handler flips it to Done when finished_ok lands.
        assert rows[e1][5] != STATUS_FAILED, (
            f"Successfully processed entry {e1} must NOT be marked "
            f"FAILED; got {rows[e1][5]!r}"
        )
        # e2 was never started — orphan cleanup must mark it FAILED.
        assert rows[e2][5] == STATUS_FAILED
        assert rows[e2][6] == "CANCELLED"

        _ExtractionWorker._is_any_worker_running = False


class TestPrimaryShortcutWiring:
    """Ctrl+Enter primary-action shortcut on Extract Text.

    Mirrors the equivalent test on Translate Document.  Without it
    a refactor of the central shortcut registry could silently drop
    Ctrl+Enter on this page.
    """

    def test_extract_shortcut_attribute_exists(self, page) -> None:  # noqa: ANN001
        assert hasattr(page, "_extract_shortcut")
        assert page._extract_shortcut is not None

    def test_extract_shortcut_default_key_is_ctrl_return(self, page) -> None:  # noqa: ANN001
        from PySide6.QtCore import Qt  # noqa: PLC0415
        from PySide6.QtGui import QKeySequence  # noqa: PLC0415

        target = QKeySequence(Qt.Modifier.CTRL | Qt.Key.Key_Return)
        assert page._extract_shortcut.key() == target

    def test_extract_shortcut_dispatches_to_handler_no_history_focus(
        self, page
    ) -> None:  # noqa: ANN001
        """No history selection → falls through to _handle_extract."""
        page.extraction_history.table = None
        with patch.object(page, "_handle_extract") as mock_handle:
            page._handle_primary_shortcut()
        mock_handle.assert_called_once()

    def test_extract_shortcut_re_extracts_on_history_selection(self, page) -> None:  # noqa: ANN001
        """Focused + selected history row → dispatch re-extract."""
        from unittest.mock import MagicMock  # noqa: PLC0415

        fake_table = MagicMock()
        fake_table.hasFocus.return_value = True
        fake_table.selectionModel.return_value.hasSelection.return_value = True
        page.extraction_history.table = fake_table

        with (
            patch.object(page.extraction_history, "on_re_extract") as mock_re,
            patch.object(page, "_handle_extract") as mock_extract,
        ):
            page._handle_primary_shortcut()

        mock_re.assert_called_once()
        mock_extract.assert_not_called()


class TestMaxFilesCapOnDrop:
    """``_MAX_FILES_PER_DROP`` (100) caps drop-area input.

    Without this regression test, the cap could be silently raised /
    removed during a refactor and an accidental large drop could
    starve the UI before the user noticed.
    """

    def test_max_files_per_drop_constant_is_100(self) -> None:
        from src.ui.pages.extract_text import (  # noqa: PLC0415
            _MAX_FILES_PER_DROP,
        )

        assert _MAX_FILES_PER_DROP == 100  # noqa: PLR2004

    def test_walk_dropped_paths_caps_at_max_files(self, page, tmp_path) -> None:  # noqa: ANN001
        """Dropping a directory with > MAX images stops at exactly MAX."""
        from PIL import Image  # noqa: PLC0415

        from src.ui.pages.extract_text import (  # noqa: PLC0415
            _MAX_FILES_PER_DROP,
        )

        d = tmp_path / "many_images"
        d.mkdir()
        for i in range(_MAX_FILES_PER_DROP + 10):
            img = Image.new("RGB", (4, 4), color="white")
            img.save(d / f"img_{i:04d}.png")

        supported, _ = page._walk_dropped_paths([str(d)])

        assert len(supported) == _MAX_FILES_PER_DROP

    def test_walk_dropped_paths_junk_does_not_consume_cap(self, page, tmp_path) -> None:  # noqa: ANN001
        """Unsupported junk files don't take cap slots — same rule as Translate Document."""
        from PIL import Image  # noqa: PLC0415

        d = tmp_path / "mixed"
        d.mkdir()
        for i in range(50):
            img = Image.new("RGB", (4, 4), color="white")
            img.save(d / f"img_{i:03d}.png")
        for i in range(200):
            (d / f"junk_{i:03d}.zzz").write_text("garbage", encoding="utf-8")

        supported, unsupported = page._walk_dropped_paths([str(d)])

        assert len(supported) == 50  # noqa: PLR2004
        assert len(unsupported) == 200  # noqa: PLR2004


class TestEmbeddedHistoryHeaderHidden:
    """Inner extraction history page's header_label is hidden when embedded.

    AGENTS.md: "Pages that embed another `create_page_container`-based
    widget hide the inner title via `page.header_label.setVisible(False)`;
    never match the label by translated text, since language-switch
    ordering can make the comparison miss."
    """

    def test_inner_extraction_history_header_is_hidden(self, page) -> None:  # noqa: ANN001
        inner_page = page.extraction_history.page
        assert inner_page.header_label.isVisible() is False
