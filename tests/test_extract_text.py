"""Unit tests for the Extract Text page and OCR worker."""

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from PySide6.QtWidgets import QApplication, QMainWindow

from src.core.ocr_engine import OCRResult
from src.ui.pages.extract_text import (
    _IMAGE_FILTER,
    ExtractTextPage,
    _ExtractionWorker,
    _write_extraction_output,
    create_extract_text_page,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def window(qapp: QApplication) -> QMainWindow:
    """Creates a QMainWindow for dialog parenting."""
    return QMainWindow()


@pytest.fixture()
def page(window: QMainWindow) -> ExtractTextPage:
    """Creates an ExtractTextPage instance."""
    return ExtractTextPage(window)


# ---------------------------------------------------------------------------
# Factory function
# ---------------------------------------------------------------------------


def test_create_extract_text_page_returns_widget(
    window: QMainWindow,
) -> None:
    """create_extract_text_page returns a QWidget."""
    widget = create_extract_text_page(window)
    assert isinstance(widget, ExtractTextPage)


# ---------------------------------------------------------------------------
# Image filter constant
# ---------------------------------------------------------------------------


def test_image_filter_contains_extensions() -> None:
    """_IMAGE_FILTER includes common image extensions."""
    assert "*.png" in _IMAGE_FILTER
    assert "*.jpg" in _IMAGE_FILTER
    assert "*.jpeg" in _IMAGE_FILTER
    assert "*.bmp" in _IMAGE_FILTER


def test_image_filter_has_all_files_section() -> None:
    """_IMAGE_FILTER includes an 'All Files' fallback."""
    assert "All Files" in _IMAGE_FILTER


# ---------------------------------------------------------------------------
# Initial UI state
# ---------------------------------------------------------------------------


def test_initial_extract_btn_disabled(page: ExtractTextPage) -> None:
    """Extract button is disabled when no files are selected."""
    assert not page.extract_btn.isEnabled()


def test_initial_view_shows_default(page: ExtractTextPage) -> None:
    """Stack shows default view (view 0) when no files selected."""
    assert page.stack.currentIndex() == 0


# ---------------------------------------------------------------------------
# File handling via _handle_files_dropped
# ---------------------------------------------------------------------------


@patch("src.ui.pages.extract_text.CustomMessageDialog.show_message")
def test_drop_valid_image_adds_to_selected(
    _mock_msg: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Valid image file is added to selected_files."""
    monkeypatch.setattr(page, "_add_file_widget", lambda _: None)
    img = tmp_path / "photo.png"
    img.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 100)
    page._handle_files_dropped([str(img)])
    assert str(img) in page.selected_files


@patch("src.ui.pages.extract_text.CustomMessageDialog.show_message")
def test_drop_non_image_rejected(
    mock_msg: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Non-image file is rejected and unsupported dialog is shown."""
    txt = tmp_path / "doc.txt"
    txt.write_text("hello")
    page._handle_files_dropped([str(txt)])
    assert page.selected_files == []
    mock_msg.assert_called_once()


@patch("src.ui.pages.extract_text.CustomMessageDialog.show_message")
def test_drop_empty_image_rejected(
    mock_msg: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Empty (0-byte) image file is rejected."""
    img = tmp_path / "empty.png"
    img.touch()
    page._handle_files_dropped([str(img)])
    assert page.selected_files == []
    mock_msg.assert_called_once()


@patch("src.ui.pages.extract_text.CustomMessageDialog.show_message")
def test_drop_duplicate_ignored(
    _mock_msg: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Adding the same file twice does not duplicate it."""
    monkeypatch.setattr(page, "_add_file_widget", lambda _: None)
    img = tmp_path / "dup.jpg"
    img.write_bytes(b"\xff\xd8\xff\xe0" + b"\x00" * 100)
    page._handle_files_dropped([str(img)])
    page._handle_files_dropped([str(img)])
    assert page.selected_files.count(str(img)) == 1


@patch("src.ui.pages.extract_text.CustomMessageDialog.show_message")
def test_drop_all_image_formats_accepted(
    _mock_msg: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """All supported image formats are accepted."""
    from src.constants import SUPPORTED_IMAGES as SI  # noqa: PLC0415

    monkeypatch.setattr(page, "_add_file_widget", lambda _: None)
    for ext in SI:
        f = tmp_path / f"test{ext}"
        f.write_bytes(b"\x00" * 10)
    page._handle_files_dropped([str(tmp_path / f"test{e}") for e in SI])
    for ext in SI:
        assert str(tmp_path / f"test{ext}") in page.selected_files


# ---------------------------------------------------------------------------
# UI state updates
# ---------------------------------------------------------------------------


def test_ui_state_after_adding_file(page: ExtractTextPage, tmp_path: Path) -> None:
    """After adding a file, extract btn enabled, stack shows files view."""
    img = tmp_path / "state.png"
    img.write_bytes(b"\x00" * 50)
    page.selected_files.append(str(img))
    page._update_ui_state()

    assert page.extract_btn.isEnabled()
    assert page.stack.currentIndex() == 1  # files view
    assert page.files_badge.text() == "1"


def test_ui_state_no_files_shows_default(
    page: ExtractTextPage,
) -> None:
    """No files shows default view, extract disabled."""
    page._update_ui_state()

    assert not page.extract_btn.isEnabled()
    assert page.stack.currentIndex() == 0  # default view


def test_ui_state_clear_all_resets_files(page: ExtractTextPage, tmp_path: Path) -> None:
    """Clear all resets files and switches to default view."""
    img = tmp_path / "clear.png"
    img.write_bytes(b"\x00" * 50)
    page.selected_files.append(str(img))

    page._handle_clear_all()

    assert page.selected_files == []
    assert not page.extract_btn.isEnabled()
    assert page.stack.currentIndex() == 0  # default view


# ---------------------------------------------------------------------------
# _on_finished — saves files immediately
# ---------------------------------------------------------------------------


@patch("src.ui.pages.extract_text.load_setting", return_value=False)
@patch("src.ui.pages.extract_text.update_extraction_status")
def test_on_finished_saves_txt_files(
    mock_update: MagicMock,
    _mock_load: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """OCR results are saved as .txt files and entry updated to Done."""
    img = tmp_path / "photo.png"
    img.write_bytes(b"\x00" * 10)
    fake_id = 42  # noqa: PLR2004
    page._on_finished([(fake_id, str(img), "Extracted text here")])

    out = tmp_path / "photo_extracted.txt"
    assert out.exists()
    assert out.read_text(encoding="utf-8") == "Extracted text here"
    mock_update.assert_called_once_with(fake_id, "Done", output_path=str(out))


@patch("src.ui.pages.extract_text.load_setting", return_value=False)
@patch("src.ui.pages.extract_text.update_extraction_status")
def test_on_finished_saves_multiple_files(
    mock_update: MagicMock,
    _mock_load: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Multiple OCR results each create a .txt file and update DB."""
    img1 = tmp_path / "a.png"
    img1.write_bytes(b"\x00" * 10)
    img2 = tmp_path / "b.jpg"
    img2.write_bytes(b"\x00" * 10)
    page._on_finished(
        [
            (1, str(img1), "Text A"),
            (2, str(img2), "Text B"),  # noqa: PLR2004
        ]
    )

    assert (tmp_path / "a_extracted.txt").read_text("utf-8") == "Text A"
    assert (tmp_path / "b_extracted.txt").read_text("utf-8") == "Text B"
    assert mock_update.call_count == 2  # noqa: PLR2004


@patch("src.ui.pages.extract_text.load_setting", return_value=False)
@patch("src.ui.pages.extract_text.update_extraction_status")
def test_on_finished_clears_worker(
    _mock_update: MagicMock,
    _mock_load: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """After finishing, worker reference is cleared."""
    page._worker = MagicMock()
    img = tmp_path / "done.png"
    img.write_bytes(b"\x00" * 10)

    page._on_finished([(1, str(img), "text")])

    assert page._worker is None


@patch("src.ui.pages.extract_text.load_setting", return_value=False)
@patch("src.ui.pages.extract_text.update_extraction_status")
def test_on_finished_write_error_marks_failed(
    mock_update: MagicMock,
    _mock_load: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """OSError during save updates entry to Failed."""
    img = tmp_path / "err.png"
    img.write_bytes(b"\x00" * 10)
    fake_id = 99  # noqa: PLR2004

    with patch("pathlib.Path.write_text", side_effect=OSError("disk full")):
        page._on_finished([(fake_id, str(img), "text")])

    mock_update.assert_called_once()
    assert mock_update.call_args[0][1] == "Failed"


# ---------------------------------------------------------------------------
# _handle_extract — OCR setup check
# ---------------------------------------------------------------------------


def test_handle_extract_no_files_is_noop(
    page: ExtractTextPage,
) -> None:
    """Extract with empty file list does nothing."""
    page.selected_files = []
    page._handle_extract()
    assert page._worker is None


@patch("src.ui.pages.extract_text.check_llm_setup", return_value=False)
@patch("src.ui.pages.extract_text.check_ocr_setup", return_value=False)
@patch(
    "src.ui.dialogs.CustomConfirmDialog.confirm",
    return_value=False,
)
def test_handle_extract_no_method_available_shows_dialog(
    mock_confirm: MagicMock,
    _mock_ocr: MagicMock,
    _mock_llm: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Extract without OCR or LLM configured shows confirmation dialog."""
    img = tmp_path / "nosetup.png"
    img.write_bytes(b"\x00" * 10)
    page.selected_files = [str(img)]

    page._handle_extract()

    mock_confirm.assert_called_once()
    assert page._worker is None


@patch("src.ui.pages.extract_text.check_llm_setup", return_value=False)
@patch("src.ui.pages.extract_text.check_ocr_setup", return_value=False)
@patch(
    "src.ui.dialogs.CustomConfirmDialog.confirm",
    return_value=True,
)
def test_handle_extract_navigates_to_settings_on_confirm(
    mock_confirm: MagicMock,
    _mock_ocr: MagicMock,
    _mock_llm: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Confirming setup dialog navigates to settings OCR tab."""
    img = tmp_path / "nav.png"
    img.write_bytes(b"\x00" * 10)
    page.selected_files = [str(img)]

    page.window_context.navigate_to_settings_tab = MagicMock()
    page._handle_extract()

    page.window_context.navigate_to_settings_tab.assert_called_once_with(
        3  # OCR tab (after General + Shortcuts + Cloud)
    )


@patch("src.ui.pages.extract_text.check_ocr_setup", return_value=True)
@patch(
    "src.ui.pages.extract_text.SourceLanguageDialog.get_selection",
    return_value=("Japanese", "", None, True),
)
@patch("src.ui.pages.extract_text.add_extraction_entry", return_value=1)
@patch("src.ui.pages.extract_text._ExtractionWorker")
def test_handle_extract_shows_dialog_and_starts_worker(  # noqa: PLR0913
    mock_worker_cls: MagicMock,
    mock_add: MagicMock,
    mock_dialog: MagicMock,
    mock_ocr: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Extract creates Pending entry, shows dialog, starts worker."""
    _settings = {
        "ocr/method": "EasyOCR",
        "extraction/method": "OCR",
        "extraction/last_output_format": ".txt",
    }
    monkeypatch.setattr(
        "src.ui.pages.extract_text.load_setting",
        lambda key, default="": _settings.get(key, default),
    )

    img = tmp_path / "method.png"
    img.write_bytes(b"\x00" * 10)
    page.selected_files = [str(img)]

    page._handle_extract()

    mock_add.assert_called_once()
    assert mock_add.call_args[1]["status"] == "Pending"
    mock_dialog.assert_called_once()
    mock_worker_cls.assert_called_once_with(
        [(1, str(img))],
        "EasyOCR",
        "Japanese",
        "OCR",
        llm_provider=None,
        llm_model=None,
    )


# ---------------------------------------------------------------------------
# _ExtractionWorker
# ---------------------------------------------------------------------------


def test_ocr_worker_emits_results(qapp: QApplication) -> None:
    """Worker emits finished_ok with (entry_id, file_path, text) tuples."""
    results_received: list[list] = []

    fake_ocr = OCRResult("Hello World", 0, 0, 100, 20, 0.95)

    with (
        patch(
            "src.ui.pages.extract_text.run_ocr",
            return_value=[fake_ocr],
        ),
        patch(
            "src.ui.pages.extract_text.merge_ocr_results",
            return_value=[fake_ocr],
        ),
    ):
        worker = _ExtractionWorker([(1, "/fake/img.png")], "TesseractOCR", "")
        worker.finished_ok.connect(results_received.append)
        worker.run()  # Run synchronously (not .start())

    assert len(results_received) == 1
    assert results_received[0][0] == (1, "/fake/img.png", "Hello World")


def test_ocr_worker_emits_progress(qapp: QApplication) -> None:
    """Worker emits progress(current, total) for each image."""
    progress_calls: list[tuple[int, int]] = []
    fake_ocr = OCRResult("Text", 0, 0, 50, 10, 0.9)

    with (
        patch(
            "src.ui.pages.extract_text.run_ocr",
            return_value=[fake_ocr],
        ),
        patch(
            "src.ui.pages.extract_text.merge_ocr_results",
            return_value=[fake_ocr],
        ),
    ):
        worker = _ExtractionWorker(
            [(1, "/fake/a.png"), (2, "/fake/b.png")], "TesseractOCR", ""
        )
        worker.progress.connect(lambda c, t: progress_calls.append((c, t)))
        worker.run()

    assert (1, 2) in progress_calls  # noqa: PLR2004
    assert (2, 2) in progress_calls  # noqa: PLR2004


@patch("src.ui.pages.extract_text.update_extraction_status")
def test_ocr_worker_import_error_marks_failed(
    mock_update: MagicMock,
    qapp: QApplication,
) -> None:
    """ImportError from run_ocr marks the task as Failed and continues."""
    results: list[object] = []

    with patch(
        "src.ui.pages.extract_text.run_ocr",
        side_effect=ImportError("tesseract not found"),
    ):
        worker = _ExtractionWorker([(1, "/fake/img.png")], "TesseractOCR", "")
        worker.finished_ok.connect(results.append)
        worker.run()

    assert len(results) == 1
    assert results[0] == []
    mock_update.assert_any_call(1, "Failed", error_message="tesseract not found")


@patch("src.ui.pages.extract_text.update_extraction_status")
def test_ocr_worker_runtime_error_marks_failed(
    mock_update: MagicMock,
    qapp: QApplication,
) -> None:
    """RuntimeError from run_ocr marks the task as Failed and continues."""
    results: list[object] = []

    with patch(
        "src.ui.pages.extract_text.run_ocr",
        side_effect=RuntimeError("engine crashed"),
    ):
        worker = _ExtractionWorker([(1, "/fake/img.png")], "TesseractOCR", "")
        worker.finished_ok.connect(results.append)
        worker.run()

    assert len(results) == 1
    assert results[0] == []
    mock_update.assert_any_call(1, "Failed", error_message="engine crashed")


@patch("src.ui.pages.extract_text.update_extraction_status")
def test_ocr_worker_value_error_marks_failed(
    mock_update: MagicMock,
    qapp: QApplication,
) -> None:
    """ValueError (AUTH_ERROR) from run_ocr marks the task as Failed."""
    results: list[object] = []

    with patch(
        "src.ui.pages.extract_text.run_ocr",
        side_effect=ValueError("AUTH_ERROR"),
    ):
        worker = _ExtractionWorker([(1, "/fake/img.png")], "TesseractOCR", "")
        worker.finished_ok.connect(results.append)
        worker.run()

    assert len(results) == 1
    assert results[0] == []
    mock_update.assert_any_call(1, "Failed", error_message="AUTH_ERROR")


@patch("src.ui.pages.extract_text.update_extraction_status")
def test_ocr_worker_generic_exception_marks_failed(
    mock_update: MagicMock,
    qapp: QApplication,
) -> None:
    """Generic exception from run_ocr marks the task as Failed."""
    results: list[object] = []

    with patch(
        "src.ui.pages.extract_text.run_ocr",
        side_effect=Exception("unexpected"),
    ):
        worker = _ExtractionWorker([(1, "/fake/img.png")], "TesseractOCR", "")
        worker.finished_ok.connect(results.append)
        worker.run()

    assert len(results) == 1
    assert results[0] == []
    mock_update.assert_any_call(1, "Failed", error_message="unexpected")


def test_ocr_worker_stop_aborts(qapp: QApplication) -> None:
    """Worker stops processing after stop() is called."""
    results: list[object] = []

    with patch(
        "src.ui.pages.extract_text.run_ocr",
        return_value=[],
    ):
        worker = _ExtractionWorker(
            [(1, "/fake/a.png"), (2, "/fake/b.png")], "TesseractOCR", ""
        )
        worker.stop()  # Stop before running
        worker.finished_ok.connect(results.append)
        worker.run()

    # finished_ok is always emitted, but with empty results when stopped
    assert len(results) == 1
    assert results[0] == []


def test_ocr_worker_empty_text_filtered(
    qapp: QApplication,
) -> None:
    """OCR results with whitespace-only text are filtered out."""
    results_received: list[list[tuple[str, str]]] = []

    fake_results = [
        OCRResult("Hello", 0, 0, 50, 10, 0.9),
        OCRResult("   ", 0, 20, 50, 10, 0.5),  # Whitespace only
        OCRResult("World", 0, 40, 50, 10, 0.9),
    ]

    with (
        patch(
            "src.ui.pages.extract_text.run_ocr",
            return_value=fake_results,
        ),
        patch(
            "src.ui.pages.extract_text.merge_ocr_results",
            return_value=fake_results,
        ),
    ):
        worker = _ExtractionWorker([(1, "/fake/img.png")], "TesseractOCR", "")
        worker.finished_ok.connect(results_received.append)
        worker.run()

    text = results_received[0][0][2]  # (entry_id, path, text)
    assert "Hello" in text
    assert "World" in text
    # Whitespace-only result should be filtered
    lines = [ln for ln in text.split("\n") if ln.strip()]
    assert len(lines) == 2  # noqa: PLR2004


def test_ocr_worker_passes_src_lang(qapp: QApplication) -> None:
    """Worker passes src_lang to run_ocr."""
    with (
        patch(
            "src.ui.pages.extract_text.run_ocr",
            return_value=[],
        ) as mock_ocr,
        patch(
            "src.ui.pages.extract_text.merge_ocr_results",
            return_value=[],
        ),
    ):
        worker = _ExtractionWorker([(1, "/fake/img.png")], "EasyOCR", "Japanese")
        worker.run()

    mock_ocr.assert_called_once_with(
        "/fake/img.png", src_lang="Japanese", method="EasyOCR"
    )


# ---------------------------------------------------------------------------
# apply_theme / apply_language
# ---------------------------------------------------------------------------


def test_apply_theme_does_not_crash(page: ExtractTextPage) -> None:
    """apply_theme runs without error."""
    page.apply_theme()


def test_apply_language_does_not_crash(
    page: ExtractTextPage,
) -> None:
    """apply_language runs without error."""
    page.apply_language()


# ---------------------------------------------------------------------------
# Window integration
# ---------------------------------------------------------------------------


def test_window_page_indices() -> None:
    """Page constants are ordered correctly."""
    from src.ui.window import (  # noqa: PLC0415
        PAGE_ABOUT,
        PAGE_DUBBING,
        PAGE_EXTRACT_TEXT,
        PAGE_GLOSSARY,
        PAGE_LIVE,
        PAGE_SETTINGS,
        PAGE_SUBTITLE,
        PAGE_TRANSLATE,
        PAGE_TRANSLATE_TEXT,
        PAGE_VOICE,
    )

    assert PAGE_TRANSLATE_TEXT == 0
    assert PAGE_TRANSLATE == 1
    assert PAGE_SUBTITLE == 2  # noqa: PLR2004
    assert PAGE_VOICE == 3  # noqa: PLR2004
    assert PAGE_DUBBING == 4  # noqa: PLR2004
    assert PAGE_LIVE == 5  # noqa: PLR2004
    assert PAGE_EXTRACT_TEXT == 6  # noqa: PLR2004
    assert PAGE_GLOSSARY == 7  # noqa: PLR2004
    assert PAGE_SETTINGS == 8  # noqa: PLR2004
    assert PAGE_ABOUT == 9  # noqa: PLR2004


def test_sidebar_keys_order() -> None:
    """Sidebar keys match the page order."""
    from src.ui.window import _SIDEBAR_KEYS  # noqa: PLC0415

    assert _SIDEBAR_KEYS[0] == "sidebar.translate_text"
    assert _SIDEBAR_KEYS[1] == "sidebar.translate_document"
    assert _SIDEBAR_KEYS[2] == "sidebar.generate_subtitle"  # noqa: PLR2004
    assert _SIDEBAR_KEYS[3] == "sidebar.generate_voice"  # noqa: PLR2004
    assert _SIDEBAR_KEYS[4] == "sidebar.dubbing"  # noqa: PLR2004
    assert _SIDEBAR_KEYS[5] == "sidebar.live"  # noqa: PLR2004
    assert _SIDEBAR_KEYS[6] == "sidebar.extract_text"  # noqa: PLR2004
    assert _SIDEBAR_KEYS[7] == "sidebar.glossary"  # noqa: PLR2004
    assert _SIDEBAR_KEYS[8] == "sidebar.settings"  # noqa: PLR2004
    assert _SIDEBAR_KEYS[9] == "sidebar.about"  # noqa: PLR2004


# ---------------------------------------------------------------------------
# _handle_files_dropped — edge cases
# ---------------------------------------------------------------------------


@patch("src.ui.pages.extract_text.CustomMessageDialog.show_message")
def test_handle_files_dropped_directory_traversal(
    mock_msg: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Dropping a directory adds all nested image files."""
    sub = tmp_path / "photos" / "vacation"
    sub.mkdir(parents=True)
    (sub / "a.png").write_bytes(b"\x00" * 10)
    (sub / "b.jpg").write_bytes(b"\x00" * 10)
    (sub / "readme.txt").write_text("not an image")

    # Stub out widget creation to prevent Qt event processing side-effects.
    monkeypatch.setattr(page, "_add_file_widget", lambda _path: None)

    page._handle_files_dropped([str(tmp_path / "photos")])

    assert len(page.selected_files) == 2  # noqa: PLR2004
    names = [Path(f).name for f in page.selected_files]
    assert "a.png" in names
    assert "b.jpg" in names
    # Unsupported dialog was shown for readme.txt
    mock_msg.assert_called_once()


@patch("src.ui.pages.extract_text.QFileDialog.getOpenFileNames")
def test_handle_files_dropped_opens_browse_on_empty(
    mock_dialog: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Empty file list triggers QFileDialog."""
    mock_dialog.return_value = ([], "")
    page._handle_files_dropped([])
    mock_dialog.assert_called_once()


@patch("src.ui.pages.extract_text.CustomMessageDialog.show_message")
def test_handle_files_dropped_shows_unsupported_dialog(
    mock_msg: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Dropping unsupported files shows a dialog listing them."""
    txt = tmp_path / "report.pdf"
    txt.write_text("not an image")

    page._handle_files_dropped([str(txt)])

    mock_msg.assert_called_once()


# ---------------------------------------------------------------------------
# _handle_remove_file
# ---------------------------------------------------------------------------


def test_handle_remove_file_updates_list_and_state(
    page: ExtractTextPage, tmp_path: Path
) -> None:
    """Removing a file updates selected_files and UI state."""
    img = tmp_path / "remove.png"
    img.write_bytes(b"\x00" * 10)
    page.selected_files = [str(img)]

    widget = MagicMock()
    page._handle_remove_file(str(img), widget)

    assert page.selected_files == []
    widget.deleteLater.assert_called_once()


def test_handle_remove_file_nonexistent_still_cleans_widget(
    page: ExtractTextPage,
) -> None:
    """Removing a file not in the list still cleans up the widget."""
    widget = MagicMock()
    page._handle_remove_file("/not/in/list.png", widget)
    widget.setParent.assert_called_once_with(None)
    widget.deleteLater.assert_called_once()


# ---------------------------------------------------------------------------
# apply_language — comprehensive
# ---------------------------------------------------------------------------


def test_apply_language_updates_all_widgets(
    page: ExtractTextPage,
) -> None:
    """apply_language updates all translatable widgets."""
    page.apply_language()

    # Verify key widgets have non-empty text
    assert len(page.extract_btn.text()) > 0
    assert len(page.clear_all_btn.text()) > 0
    assert len(page.section_label.text()) > 0
    # Drop area supported label is updated (non-empty)
    label_text = page.drop_area.supported_label.text()
    assert len(label_text) > 0


# ---------------------------------------------------------------------------
# _write_extraction_output
# ---------------------------------------------------------------------------


def test_write_extraction_output_txt(tmp_path: Path) -> None:
    """Plain text content is written to .txt files as UTF-8."""
    out = tmp_path / "result.txt"
    _write_extraction_output(out, "Hello, World!")
    assert out.read_text(encoding="utf-8") == "Hello, World!"


def test_write_extraction_output_multiline_txt(tmp_path: Path) -> None:
    """Multi-line text is preserved verbatim."""
    text = "Line 1\nLine 2\nLine 3"
    out = tmp_path / "result.txt"
    _write_extraction_output(out, text)
    assert out.read_text(encoding="utf-8") == text


def test_write_extraction_output_docx_creates_file(
    tmp_path: Path,
) -> None:
    """Writing .docx creates a valid binary file."""
    out = tmp_path / "result.docx"
    _write_extraction_output(out, "Some text\nMore text")
    assert out.exists()
    assert out.stat().st_size > 0


def test_write_extraction_output_unicode(tmp_path: Path) -> None:
    """Unicode characters are stored correctly."""
    text = "Xin chào 你好 مرحبا"
    out = tmp_path / "result.txt"
    _write_extraction_output(out, text)
    assert out.read_text(encoding="utf-8") == text


def test_write_extraction_output_empty_string(tmp_path: Path) -> None:
    """Empty text is written without error."""
    out = tmp_path / "empty.txt"
    _write_extraction_output(out, "")
    assert out.exists()
    assert out.read_text(encoding="utf-8") == ""


# ---------------------------------------------------------------------------
# _on_finished — auto_remove path
# ---------------------------------------------------------------------------


@patch("src.ui.pages.extract_text.load_setting", return_value=True)
@patch("src.ui.pages.extract_text.delete_extraction_entry")
def test_on_finished_auto_remove_deletes_entry(
    mock_delete: MagicMock,
    _mock_load: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """When auto_remove is enabled, finished entries are deleted from DB."""
    img = tmp_path / "autoremove.png"
    img.write_bytes(b"\x00" * 10)
    fake_id = 77  # noqa: PLR2004

    page._on_finished([(fake_id, str(img), "Extracted text")])

    mock_delete.assert_called_once_with(fake_id)


@patch("src.ui.pages.extract_text.load_setting", return_value=True)
@patch("src.ui.pages.extract_text.delete_extraction_entry")
@patch("src.ui.pages.extract_text.update_extraction_status")
def test_on_finished_auto_remove_does_not_call_update(
    mock_update: MagicMock,
    mock_delete: MagicMock,
    _mock_load: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """With auto_remove=True, update_extraction_status is NOT called."""
    img = tmp_path / "skip_update.png"
    img.write_bytes(b"\x00" * 10)

    page._on_finished([(1, str(img), "text")])

    mock_update.assert_not_called()
    mock_delete.assert_called_once()


# ---------------------------------------------------------------------------
# _on_finished — fallback when _output_format is missing
# ---------------------------------------------------------------------------


@patch("src.ui.pages.extract_text.update_extraction_status")
@patch("src.ui.pages.extract_text._write_extraction_output")
@patch("src.ui.pages.extract_text.CustomMessageDialog.show_message")
def test_on_finished_default_output_format_is_txt(
    _mock_msg: MagicMock,
    mock_write: MagicMock,
    _mock_update: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Default _output_format (.txt from __init__) produces .txt output."""
    img = tmp_path / "test.png"
    img.write_bytes(b"\x00" * 10)

    from src.core.database import add_extraction_entry  # noqa: PLC0415

    fake_id = add_extraction_entry("test.png", 10, str(img), "", "Pending")
    page._pending_tasks = [(fake_id, str(img))]
    page._worker = MagicMock()

    # _output_format defaults to ".txt" from __init__ — no _start_worker call
    page._on_finished([(fake_id, str(img), "Extracted text")])

    # _write_extraction_output should have been called with .txt extension
    assert mock_write.called
    output_path = mock_write.call_args[0][0]
    assert str(output_path).endswith(".txt")


# ---------------------------------------------------------------------------
# _ExtractionWorker — LLM extraction path
# ---------------------------------------------------------------------------


def test_extraction_worker_llm_path(qapp: QApplication) -> None:
    """Worker with LLM method calls _extract_with_llm instead of OCR."""
    results_received: list[list] = []

    with patch(
        "src.core.llm_engine.extract_image_text",
        return_value="LLM extracted text",
    ):
        worker = _ExtractionWorker(
            [(1, "/fake/img.png")],
            "TesseractOCR",
            "",
            extract_method="LLM",
        )
        worker.finished_ok.connect(results_received.append)
        worker.run()

    assert len(results_received) == 1
    entry_id, path, text = results_received[0][0]
    assert text == "LLM extracted text"


@patch("src.ui.pages.extract_text.update_extraction_status")
def test_extraction_worker_llm_value_error_marks_failed(
    mock_update: MagicMock,
    qapp: QApplication,
) -> None:
    """LLM ValueError (e.g. AUTH_ERROR) marks the task as Failed."""
    results: list[object] = []

    with patch(
        "src.core.llm_engine.extract_image_text",
        side_effect=ValueError("AUTH_ERROR"),
    ):
        worker = _ExtractionWorker(
            [(1, "/fake/img.png")],
            "TesseractOCR",
            "",
            extract_method="LLM",
        )
        worker.finished_ok.connect(results.append)
        worker.run()

    assert len(results) == 1
    assert results[0] == []
    mock_update.assert_any_call(1, "Failed", error_message="AUTH_ERROR")


@patch("src.ui.pages.extract_text.update_extraction_status")
def test_extraction_worker_llm_generic_exception_marks_failed(
    mock_update: MagicMock,
    qapp: QApplication,
) -> None:
    """LLM generic exception marks the task as Failed."""
    results: list[object] = []

    with patch(
        "src.core.llm_engine.extract_image_text",
        side_effect=RuntimeError("API timeout"),
    ):
        worker = _ExtractionWorker(
            [(1, "/fake/img.png")],
            "TesseractOCR",
            "",
            extract_method="LLM",
        )
        worker.finished_ok.connect(results.append)
        worker.run()

    assert len(results) == 1
    assert results[0] == []
    mock_update.assert_any_call(1, "Failed", error_message="API timeout")


# ---------------------------------------------------------------------------
# extract_image_text — dispatch
# ---------------------------------------------------------------------------


def test_extract_image_text_gemini_dispatch() -> None:
    """extract_image_text dispatches to Gemini when method is Gemini."""
    from src.core.llm_engine import extract_image_text  # noqa: PLC0415

    with (
        patch(
            "src.core.llm_engine._resolve_provider_model",
            return_value=("Gemini", "gemini-3-flash-preview"),
        ),
        patch(
            "src.core.llm_engine._extract_text_gemini",
            return_value="gemini result",
        ) as mock_gemini,
    ):
        result = extract_image_text("/fake/img.png")

    assert result == "gemini result"
    # extract_image_text now passes the selected model id along to the
    # provider dispatcher (new `model` positional argument).
    mock_gemini.assert_called_once()
    assert mock_gemini.call_args.args[0] == "/fake/img.png"


def test_extract_image_text_custom_dispatch() -> None:
    """extract_image_text dispatches to Custom when method is Custom."""
    from src.core.llm_engine import extract_image_text  # noqa: PLC0415

    with (
        patch(
            "src.core.llm_engine._resolve_provider_model",
            return_value=("Custom", "gpt-4o"),
        ),
        patch(
            "src.core.llm_engine._extract_text_custom",
            return_value="custom result",
        ) as mock_custom,
    ):
        result = extract_image_text("/fake/img.png")

    assert result == "custom result"
    mock_custom.assert_called_once()
    assert mock_custom.call_args.args[0] == "/fake/img.png"


def test_extract_image_text_unknown_method_returns_empty() -> None:
    """extract_image_text returns empty string when resolver returns an unknown provider."""
    from src.core.llm_engine import extract_image_text  # noqa: PLC0415

    # _resolve_provider_model now drives dispatch; bypass it with a direct
    # stub that reports an unrecognised provider so we exercise the
    # ``return ""`` branch.
    with patch(
        "src.core.llm_engine._resolve_provider_model",
        return_value=("UnknownProvider", ""),
    ):
        result = extract_image_text("/fake/img.png")

    assert result == ""


# ---------------------------------------------------------------------------
# _extract_text_custom — missing credentials
# ---------------------------------------------------------------------------


def test_extract_text_custom_missing_credentials() -> None:
    """_extract_text_custom raises AUTH_ERROR when credentials missing."""
    from src.core.llm_engine import _extract_text_custom  # noqa: PLC0415

    with (
        patch("src.core.llm_engine._config.load_setting", return_value=""),
        pytest.raises(ValueError, match="AUTH_ERROR"),
    ):
        _extract_text_custom("/fake/img.png")


# ---------------------------------------------------------------------------
# _handle_re_extract — requirements and dialog
# ---------------------------------------------------------------------------


@patch("src.ui.pages.extract_text.check_llm_setup", return_value=False)
@patch("src.ui.pages.extract_text.check_ocr_setup", return_value=False)
@patch(
    "src.ui.dialogs.CustomConfirmDialog.confirm",
    return_value=False,
)
def test_handle_re_extract_no_method_shows_dialog(
    mock_confirm: MagicMock,
    _mock_ocr: MagicMock,
    _mock_llm: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Re-extract without OCR or LLM shows setup-required dialog."""
    page._handle_re_extract([(1, str(tmp_path / "img.png"))])
    mock_confirm.assert_called_once()


@patch("src.ui.pages.extract_text.check_ocr_setup", return_value=True)
@patch(
    "src.ui.pages.extract_text.SourceLanguageDialog.get_selection",
    return_value=("", "", None, False),
)
def test_handle_re_extract_dialog_cancelled(
    mock_dialog: MagicMock,
    _mock_ocr: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Re-extract with cancelled dialog does not start worker."""
    page._handle_re_extract([(1, str(tmp_path / "img.png"))])
    mock_dialog.assert_called_once()
    assert page._worker is None


# ---------------------------------------------------------------------------
# _handle_extract — dialog cancelled
# ---------------------------------------------------------------------------


@patch("src.ui.pages.extract_text.check_ocr_setup", return_value=True)
@patch(
    "src.ui.pages.extract_text.SourceLanguageDialog.get_selection",
    return_value=("", "", None, False),
)
def test_handle_extract_dialog_cancelled_no_worker(
    _mock_dialog: MagicMock,
    _mock_ocr: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
) -> None:
    """Cancelling the source language dialog does not start a worker."""
    img = tmp_path / "cancel.png"
    img.write_bytes(b"\x00" * 10)
    page.selected_files = [str(img)]

    page._handle_extract()

    assert page._worker is None


# ---------------------------------------------------------------------------
# _extract_text_gemini — mocked API round-trip
# ---------------------------------------------------------------------------


def test_extract_text_gemini_success() -> None:
    """_extract_text_gemini returns extracted text on successful API call."""
    import json  # noqa: PLC0415
    from types import SimpleNamespace  # noqa: PLC0415

    from src.core.llm_engine import _extract_text_gemini  # noqa: PLC0415

    fake_client = MagicMock()
    fake_client.models.generate_content.return_value = SimpleNamespace(
        text=json.dumps({"text": "Hello World"}),
    )

    with (
        patch(
            "src.core.llm_engine._config.load_setting",
            side_effect=lambda k, d="": {
                "llm/gemini_api_key": "test-api-key",
                "llm/gemini_model": "gemini-2.5-flash",
            }.get(k, d),
        ),
        patch(
            "src.core.llm_engine._build_gemini_client",
            return_value=fake_client,
        ),
        patch("pathlib.Path.read_bytes", return_value=b"\x89PNG fake"),
    ):
        result = _extract_text_gemini("/fake/img.png")

    assert result == "Hello World"


def test_extract_text_gemini_api_error_raises() -> None:
    """_extract_text_gemini raises ValueError on API error."""
    from google.genai import errors  # noqa: PLC0415

    from src.core.llm_engine import _extract_text_gemini  # noqa: PLC0415

    fake_client = MagicMock()
    fake_client.models.generate_content.side_effect = errors.APIError(
        code=401,
        response_json={"error": {"message": "Unauthorized", "code": 401}},
        response=None,
    )

    with (
        patch(
            "src.core.llm_engine._config.load_setting",
            side_effect=lambda k, d="": {
                "llm/gemini_api_key": "bad-key",
                "llm/gemini_model": "gemini-2.5-flash",
            }.get(k, d),
        ),
        patch(
            "src.core.llm_engine._build_gemini_client",
            return_value=fake_client,
        ),
        patch("pathlib.Path.read_bytes", return_value=b"\x89PNG"),
        pytest.raises(ValueError, match="AUTH_ERROR"),
    ):
        _extract_text_gemini("/fake/img.png")


# ---------------------------------------------------------------------------
# _extract_text_custom — mocked API round-trip
# ---------------------------------------------------------------------------


def test_extract_text_custom_success() -> None:
    """_extract_text_custom returns extracted text on successful API call."""
    import json  # noqa: PLC0415
    from types import SimpleNamespace  # noqa: PLC0415

    from src.core.llm_engine import _extract_text_custom  # noqa: PLC0415

    fake_client = MagicMock()
    fake_client.chat.completions.create.return_value = SimpleNamespace(
        choices=[
            SimpleNamespace(
                message=SimpleNamespace(
                    content=json.dumps({"text": "Extracted by custom LLM"}),
                ),
            ),
        ],
    )
    # Vision calls go through ``client.with_options(timeout=...)`` so the
    # ``with_options`` chain must return the same mock (otherwise the
    # ``chat.completions.create`` mock above never fires).
    fake_client.with_options.return_value = fake_client

    with (
        patch(
            "src.core.llm_engine._resolve_custom_config",
            return_value=("sk-key", "gpt-4o", "https://api.example.com/v1"),
        ),
        patch(
            "src.core.llm_engine._build_openai_client",
            return_value=fake_client,
        ),
        patch("pathlib.Path.read_bytes", return_value=b"\x89PNG"),
    ):
        result = _extract_text_custom("/fake/img.png")

    assert result == "Extracted by custom LLM"


# ---------------------------------------------------------------------------
# _EXTRACT_TEXT_PROMPT — exists and is non-empty
# ---------------------------------------------------------------------------


def test_extract_text_prompt_is_valid() -> None:
    """_EXTRACT_TEXT_PROMPT is a non-empty string with key instructions."""
    from src.core.llm_engine import _EXTRACT_TEXT_PROMPT  # noqa: PLC0415

    assert isinstance(_EXTRACT_TEXT_PROMPT, str)
    assert len(_EXTRACT_TEXT_PROMPT) > 0
    assert "extract" in _EXTRACT_TEXT_PROMPT.lower()
    assert "translate" in _EXTRACT_TEXT_PROMPT.lower()  # "Do NOT translate"
    assert "JSON" in _EXTRACT_TEXT_PROMPT


# ---------------------------------------------------------------------------
# _handle_extract — all add_extraction_entry return None → no worker
# ---------------------------------------------------------------------------


@patch("src.ui.pages.extract_text.check_ocr_setup", return_value=True)
@patch(
    "src.ui.pages.extract_text.SourceLanguageDialog.get_selection",
    return_value=("", "", None, True),
)
@patch("src.ui.pages.extract_text.add_extraction_entry", return_value=None)
@patch("src.ui.pages.extract_text._ExtractionWorker")
def test_handle_extract_all_entries_fail_no_worker(
    mock_worker_cls: MagicMock,
    _mock_add: MagicMock,
    _mock_dialog: MagicMock,
    _mock_ocr: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """When all DB entries fail to create, worker is not started."""
    monkeypatch.setattr(
        "src.ui.pages.extract_text.load_setting",
        lambda key, default="": default,
    )
    img = tmp_path / "fail.png"
    img.write_bytes(b"\x00" * 10)
    page.selected_files = [str(img)]

    page._handle_extract()

    mock_worker_cls.assert_not_called()


# ---------------------------------------------------------------------------
# _handle_files_dropped — mixed valid + invalid files
# ---------------------------------------------------------------------------


@patch("src.ui.pages.extract_text.CustomMessageDialog.show_message")
def test_handle_files_dropped_mixed_valid_and_invalid(
    mock_msg: MagicMock,
    page: ExtractTextPage,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Dropping both valid images and unsupported files adds valid ones only."""
    monkeypatch.setattr(page, "_add_file_widget", lambda _: None)

    img = tmp_path / "photo.png"
    img.write_bytes(b"\x00" * 10)
    txt = tmp_path / "notes.txt"
    txt.write_text("hello")

    page._handle_files_dropped([str(img), str(txt)])

    # Valid image was added
    assert len(page.selected_files) == 1
    assert page.selected_files[0] == str(img)
    # Unsupported dialog shown for .txt
    mock_msg.assert_called_once()


# ---------------------------------------------------------------------------
# _extract_text_gemini — non-vision model fallback
# ---------------------------------------------------------------------------


def test_extract_text_gemini_non_vision_model_falls_back() -> None:
    """Non-vision model name is replaced with default vision model."""
    import json  # noqa: PLC0415
    from types import SimpleNamespace  # noqa: PLC0415

    from src.core.llm_engine import (  # noqa: PLC0415
        DEFAULT_GEMINI_MODEL,
        _extract_text_gemini,
    )

    fake_client = MagicMock()
    fake_client.models.generate_content.return_value = SimpleNamespace(
        text=json.dumps({"text": "result"}),
    )

    with (
        patch(
            "src.core.llm_engine._config.load_setting",
            side_effect=lambda k, d="": {
                "llm/gemini_api_key": "test-key",
                "llm/gemini_model": "text-only-model",  # NOT a vision model
            }.get(k, d),
        ),
        patch(
            "src.core.llm_engine._build_gemini_client",
            return_value=fake_client,
        ),
        patch("pathlib.Path.read_bytes", return_value=b"\x89PNG"),
    ):
        _extract_text_gemini("/fake/img.png")

    # Verify the SDK call used the default vision model, not "text-only-model".
    fake_client.models.generate_content.assert_called_once()
    assert (
        fake_client.models.generate_content.call_args.kwargs["model"]
        == DEFAULT_GEMINI_MODEL
    )


# ---------------------------------------------------------------------------
# _write_extraction_output — additional edge-case tests
# ---------------------------------------------------------------------------


def test_write_extraction_output_docx_paragraphs(tmp_path: Path) -> None:
    """DOCX output creates paragraphs matching each newline-delimited line."""
    out = tmp_path / "paragraphs.docx"
    _write_extraction_output(out, "Line one\nLine two\nLine three")
    assert out.exists()

    from docx import Document  # noqa: PLC0415

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert texts == ["Line one", "Line two", "Line three"]


def test_write_extraction_output_docx_empty_text(tmp_path: Path) -> None:
    """DOCX output with empty string creates a file with a single empty paragraph."""
    out = tmp_path / "empty.docx"
    _write_extraction_output(out, "")
    assert out.exists()

    from docx import Document  # noqa: PLC0415

    doc = Document(str(out))
    # Empty string split by \n yields [""], so one paragraph
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == ""


def test_write_extraction_output_docx_unicode(tmp_path: Path) -> None:
    """Non-ASCII text (CJK, accented, emoji) is preserved in .docx output."""
    text = "日本語テスト\nrésumé\n🌍"
    out = tmp_path / "unicode.docx"
    _write_extraction_output(out, text)

    from docx import Document  # noqa: PLC0415

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert texts == ["日本語テスト", "résumé", "🌍"]


def test_write_extraction_output_txt_newlines_preserved(tmp_path: Path) -> None:
    """Multiple consecutive newlines in text are preserved in .txt output."""
    text = "First\n\n\nFourth"
    out = tmp_path / "newlines.txt"
    _write_extraction_output(out, text)
    assert out.read_text(encoding="utf-8") == text
